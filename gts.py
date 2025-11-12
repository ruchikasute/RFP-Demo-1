import streamlit as st
from openai import AzureOpenAI
import os, io
from docx import Document
from PyPDF2 import PdfReader
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import re

def detect_client_name_from_text(text: str) -> str:
    """
    Smarter detection of client name from RFP or SOW text.
    Scans for patterns like 'Client:', 'Prepared for', 'RFP from', 'Issued by', etc.
    Falls back gracefully if not found.
    """
    patterns = [
        r"(?i)\bclient\s*(?:name)?\s*[:\-]\s*([A-Za-z0-9&,\.\s]+)",
        r"(?i)\bprepared\s*for\s*([A-Za-z0-9&,\.\s]+)",
        r"(?i)\bproposal\s*(?:for|to)\s*([A-Za-z0-9&,\.\s]+)",
        r"(?i)\brfp\s*(?:from|by|for)\s*([A-Za-z0-9&,\.\s]+)",
        r"(?i)\bissued\s*(?:by|to)\s*([A-Za-z0-9&,\.\s]+)",
        r"(?i)\bsubmitted\s*(?:by|to)\s*([A-Za-z0-9&,\.\s]+)",
        r"(?i)\borganization\s*[:\-]\s*([A-Za-z0-9&,\.\s]+)",
    ]

    for pat in patterns:
        match = re.search(pat, text)
        if match:
            name = match.group(1).strip()
            # clean any trailing words like "Limited", "LLC", etc.
            name = re.sub(r"\s+(Limited|Ltd|LLC|Company|Inc\.?)\b.*", r" \1", name, flags=re.I)
            # remove extra newlines or dots
            name = re.sub(r"[\n\r]+", " ", name).strip(" .")
            # cap first letters
            return name.title()

    return "Client"



def insert_formatted_text(doc, placeholder, raw_text):
    """
    Smart parser: handles any LLM output format (Markdown, numbered, or plain text)
    Converts to clean formatted Word content (Headings, bullets, tables, paragraphs).
    """

    def set_cell_shading(cell, fill_color):
        """Add shading to a table cell."""
        tc_pr = cell._element.tcPr
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)
        tc_pr.append(shd)

    def insert_styled_table(parent, headers, rows):
        """Insert styled table with blue header."""
        table = parent.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h.strip()
            set_cell_shading(hdr_cells[i], "008FD3")
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        for r, row_data in enumerate(rows):
            cells = table.rows[r + 1].cells
            for c, val in enumerate(row_data):
                cells[c].text = str(val).strip()
                set_cell_shading(cells[c], "E7EEF7")
        return table

    def add_heading(doc, text, level):
        """Add heading with correct style and spacing."""
        p = doc.add_paragraph(text.strip(), style=f"Heading {level}")
        p.paragraph_format.space_after = Pt(4)
        return p

    # --- Find placeholder ---
    inserted = False
    for para in doc.paragraphs:
        if placeholder in para.text:
            inserted = True
            parent = para._element.getparent()
            idx = parent.index(para._element)
            parent.remove(para._element)

            lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
            new_elements = []
            i = 0

            while i < len(lines):
                line = lines[i]

                # --- Markdown Headings ---
                if re.match(r"^#{1,3}\s", line):
                    level = line.count("#")
                    heading_text = re.sub(r"^#+\s*", "", line)
                    new_elements.append(add_heading(doc, heading_text, level)._element)
                    i += 1
                    continue

                # --- Numbered headings (1., 2., etc.) ---
                if re.match(r"^\d+\.\s+[A-Z]", line):
                    heading_text = re.sub(r"^\d+\.\s*", "", line)
                    new_elements.append(add_heading(doc, heading_text, 1)._element)
                    i += 1
                    continue

                # --- Bold-only headings ---
                if re.match(r"^\*\*[^\*]+\*\*$", line):
                    heading_text = re.sub(r"\*\*", "", line)
                    new_elements.append(add_heading(doc, heading_text, 2)._element)
                    i += 1
                    continue

                # --- Tables ---
                if line.startswith("|") and "|" in line:
                    table_lines = []
                    while i < len(lines) and lines[i].startswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    headers = [h.strip("* ") for h in table_lines[0].strip("|").split("|")]
                    rows = [r.strip("|").split("|") for r in table_lines[2:]]
                    table = insert_styled_table(doc, headers, rows)
                    new_elements.append(table._element)
                    continue

                # --- Bullets (dash or dot) ---
                if line.startswith("- ") or line.startswith("• "):
                    bullet_text = re.sub(r"^[-•]\s*", "", line)
                    p = doc.add_paragraph(bullet_text, style="List Bullet 2")
                    p.paragraph_format.left_indent = Pt(18)
                    p.paragraph_format.space_after = Pt(2)
                    new_elements.append(p._element)
                    i += 1
                    continue

                # --- Regular text ---
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.2
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                new_elements.append(p._element)
                i += 1

            # insert processed elements
            for el in reversed(new_elements):
                parent.insert(idx, el)
            break

    if not inserted:
        st.warning(f"⚠️ Placeholder {placeholder} not found — appending content at end.")
        doc.add_paragraph(raw_text)


# ============================================================
# Helper: Extract text from uploaded document
# ============================================================

def extract_text_from_file(uploaded_file):
    """Extract text from uploaded PDF or DOCX."""
    text = ""
    if uploaded_file.name.endswith(".pdf"):
        pdf = PdfReader(uploaded_file)
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    elif uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        st.warning("⚠️ Unsupported file type. Please upload a PDF or DOCX.")
    return text.strip()
from docx.shared import Inches

def insert_sustainability_section(doc, image_top=None, image_bottom=None):
    """
    Inserts a sustainability and EcoVadis section before 'Project Scope' section.
    Ensures correct order: Top image → Text → Bottom image.
    """
    from docx.shared import Inches, Pt
    import re, os

    sustainability_text = """
At Crave InfoTech, we recognize the importance of sustainability in fostering long-term growth, innovation, and responsibility. 
As a leading player in the IT sector, we are committed to integrating sustainable practices into every aspect of our business. 
Our approach to sustainability is guided by the principles of environmental stewardship, social responsibility, and economic viability.

**Ecovadis Rating for Crave InfoTech**
Crave InfoTech maintains a strong commitment to sustainability excellence, as reflected in our EcoVadis rating. 
This recognition underscores our focus on ethical business practices, sustainable procurement, and corporate responsibility across all engagements.
    """

    # 🔍 Find insertion point: before "Project Scope"
    target_para = None
    for para in doc.paragraphs:
        if re.search(r"\bProject\s+Scope\b", para.text, re.IGNORECASE):
            target_para = para
            break

    if target_para:
        parent = target_para._element.getparent()
        idx = parent.index(target_para._element)

        # --- Create paragraph for top image, if exists ---
        if image_top and os.path.exists(image_top):
            p_top = doc.add_paragraph()
            run = p_top.add_run()
            run.add_picture(image_top, width=Inches(6))
            parent.insert(idx, p_top._element)
            idx += 1  # move index forward

        # --- Sustainability text ---
        p_text = doc.add_paragraph(sustainability_text.strip())
        p_text.paragraph_format.space_after = Pt(6)
        for run in p_text.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)
        parent.insert(idx, p_text._element)
        idx += 1

        # --- Bottom image, if exists ---
        if image_bottom and os.path.exists(image_bottom):
            p_bottom = doc.add_paragraph()
            run = p_bottom.add_run()
            run.add_picture(image_bottom, width=Inches(3))
            parent.insert(idx, p_bottom._element)

        st.success("🌱 Sustainability section (with EcoVadis images) inserted before 'Project Scope'.")
    else:
        st.warning("⚠️ 'Project Scope' not found — appending sustainability section at document end.")
        doc.add_page_break()
        doc.add_paragraph(sustainability_text.strip())


# ============================================================
# Helper: Generate GTS RFP/SOW Response
# ============================================================

def generate_gts_sow(client, model_name,  reference_text,client_name):
    """Generate a detailed SOW response document using Azure OpenAI."""
    today = datetime.now().strftime("%d %B %Y")
    template_path = "Template/GTS_Template.docx"

    # --- Load or create Word template ---
    if os.path.exists(template_path):
        doc = Document(template_path)
        st.info("📄 Using GTS Word template.")
    else:
        st.warning("⚠️ Template not found. Creating blank document.")
        doc = Document()
        doc.add_paragraph("<<CONTENT START>>")

    # Pre-clean uploaded RFP text to remove proposal-style headers

    reference_text = re.sub(r"(?i)(proposal\s+for\s+sap\s+.*?)(\n|$)", "", reference_text)
    reference_text = re.sub(r"(?i)(^|\n)\s*1\.\s*proposal\s+for\s+.*", "", reference_text)
    reference_text = re.sub(r"(?i)sap\s+gts\s+processes\s+and\s+enhancements", "", reference_text)

    # --- Build LLM prompt ---
    prompt = f"""
You are a Senior SAP GTS consultant from Crave InfoTech preparing a professional
RFP Response / Statement of Work for client.

Below is the input RFP or reference document provided by the client or user.
Use it to infer project context, objectives, requirements, and tone.

REFERENCE DOCUMENT:
{reference_text}

Now, generate a **comprehensive and polished proposal** structured with the following sections:

Introduction - 
   The Introduction should open the document with a concise, narrative-style overview of the project’s purpose, primary objectives, and key technologies involved.  
    Expand this into 3–4 rich paragraphs that clearly connect the business objectives to Crave InfoTech’s capabilities. Describe how Crave leverages its SAP Build Process Automation (SBPA) expertise, domain understanding, and delivery accelerators to help the client achieve measurable efficiency, compliance, and scalability improvements.  
 
   Focus on how Crave InfoTech’s expertise in SAP Build Process Automation (SBPA), domain knowledge, and implementation experience will enable the client to achieve accuracy, compliance, and operational efficiency.  
   Use only the detected client name `{client_name}` where relevant, and avoid referring to unrelated organizations or prior examples (e.g., Eli Lilly, GTS Enhancements).  
   Maintain a formal and consultative tone that naturally positions Crave InfoTech as the trusted delivery partner for the initiative.  
Project Scope
    2.1 in-scope items with clarity and structure
    2.2. Solution Architecture  
    - Describe proposed SAP GTS architecture, integration, and key components
    2.3.Prerequisites and Key Assumptions  
   - Outline key preconditions and technical assumptions
    2.4.Out of Scope  
   - Clearly list all out-of-scope functionalities
    2.5.Project Documents Deliverables  
   - Enumerate Documents deliverables
Bill of Materials (BOM)  
   - Mention major components, licenses, and tools (if applicable) in tabular format
Responsibility Matrix : This section should have 2 parts
   - Define RACI (Crave / Client) responsibilities in tabular format - The table must include the following columns:
   *Task* | *Crave InfoTech (R/A/C/I)* | *client (R/A/C/I)*  

   - **Part 2 – Roles & Responsibilities Table:**  
        Below the RACI table, provide a second table titled “Roles & Responsibilities” with two columns:  
        1. *Role*  
        2. *Key Responsibilities*  
Project Delivery Approach 
   - Project Organization Structure
   - Project Resource Planning
        - Begin with a short paragraph describing how Crave’s Delivery and Project Managers identify and allocate resources during project initiation.  
        - Then generate a detailed **Project Resource Planning Table** that lists project phases vs. indicative roles and resource involvement.  
        - Title the table as **“Exhibit: Project Resource Planning”**, and include the following columns:  
            - *Project Phase*  
            - *Functional Consultant*  
            - *Technical Consultant*  
            - *Other Roles (Developer, Tester, PM, ABAP, Architect, etc.)*  
   - The rows should include realistic phases such as:  
       - Project Preparation  
       - Business Blueprint / Detailed Design  
       - Realization – Development  
       - Unit and Integration Testing  
       - Documentation and User Training  
       - Cutover and Go-Live  
       - Post Go-Live Support  
   - Fill each cell with “Yes” or short role mentions (e.g., “PM, Functional Consultant, and Developer”) as per typical Crave-style staffing exhibits.  
   - Implementation Methodology
   - Communication Plan 
       This section should describe how Crave InfoTech and the Client will communicate and manage issues throughout the project lifecycle.  
   
        **Part 1 – Interaction Table:**  
        Create a table titled “Communication Schedule” with columns such as *Interaction*, *Frequency*, and *Purpose* to outline daily, weekly, and monthly touchpoints.  

        **Part 2 – Issue Management and Escalation Process:**  
        Add a second table titled “Issue Management and Escalation Process” that includes columns like *Task*, *Timescale*, *Resposibility*.  

        **Part 3 – Issue Classification Table:**  
        Include a third table titled “Issue Classification” to categorize issues by severity (e.g., Low, Serious, Critical) with Definition, Reporting Process and Solution Responsible. 

        Part 4 - Escalation Process
        Include fourth table titles as "Escalation Process"  to categorize escalation handling across roles and governance levels.  
        - The table should have the following columns: *Issue Type*, *Escalation Point*, *Escalation Criteria*.  
        - Populate the table with realistic entries such as:  
            - Governance Role (Project Core Group)  
            - Project Delivery Manager  
            - Crave Technology Project Manager  
            - Crave Technology Delivery Manager  
        - Each row should include escalation criteria (e.g., “If plan to resolve the issue is not outlined within 48 hrs”, “Weekly checkpoints”, “Quality issues”, “Unresolved delivery issue”, etc.).  


Timelines  
   6.1 Delivery Timeliness
   6.2 Efforts and Resource Allocation
   6.3 Commercials
   6.4 Payment Terms - - List typical payment milestones in tabular format
Sign-Off  
   - Add formal sign-off language
Other Assumptions  
   - Highlight additional assumptions for clarity
    8.1. Dependency
    8.2. Limitations  
   - Mention project or technical limitations
    8.3. General Provisions  
   - Include standard terms and closing remarks

Each section must be clearly titled (Heading 1) and contain 1–2 detailed paragraphs written
in Crave Infotech’s professional proposal tone.

Ensure:
- Content flows logically and professionally.
- No bullet formatting unless it improves clarity.
- Tone: confident, formal, consultative.
- Personalize context for client.
- The document **must begin directly with the section heading "1. Introduction"**, without any preamble, title, or summary lines like “Proposal for …”.
"""

    # --- Call Azure LLM ---
    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4
        )
        sow_text = response.choices[0].message.content.strip()
        # --- 🧹 Clean unwanted headers ---
        cleanup_patterns = [
            r"(?i)^crave\s+infotech\s+proposal.*\n?",
            r"(?i)^sap\s+gts\s+implementation.*\n?",
            r"(?i)^proposal\s+for.*\n?",
            r"(?i)^client\s*:\s*.*\n?",
            r"(?i)^date\s*:\s*.*\n?",
            r"(?i)^---+\s*\n?",
        ]
        for pattern in cleanup_patterns:
            sow_text = re.sub(pattern, "", sow_text).strip()

        # --- Force start from Introduction ---
        intro_match = re.search(r"(?i)(^|\n)(\d+\.\s*)?introduction", sow_text)
        if intro_match:
            sow_text = sow_text[intro_match.start():].strip()
    except Exception as e:
        st.error(f"⚠️ Error calling model: {e}")
        sow_text = "Error generating document content."


    # --- Insert formatted content instead of plain text ---
    insert_formatted_text(doc, "<<CONTENT START>>", sow_text)

    # --- Add client name and metadata at top ---
    doc.add_paragraph()

    insert_sustainability_section(
        doc,
        image_top="Images/Crave Awards.png",   # optional top banner
        image_bottom="Images/Sustainability.png"       # optional EcoVadis badge
    )
    # --- Save to memory ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ============================================================
# Streamlit UI
# ============================================================

def main():
    st.title("🌐 GTS — SOW Generator")

    uploaded_file = st.file_uploader(
        " ",
        type=["pdf", "docx"],
        key="rfp_uploader",
        help="Upload the client RFP or reference document (PDF/DOCX).",
        label_visibility="collapsed"
    )

    reference_text = ""
    if uploaded_file:
        reference_text = extract_text_from_file(uploaded_file)
        st.success(f"✅ Extracted text from `{uploaded_file.name}` ({len(reference_text.split())} words)")
        client_name = detect_client_name_from_text(reference_text)
        st.info(f"📌 Detected Client Name: **{client_name}**")

    # Azure setup
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )
    model_name = "codetest"

    if st.button("⚡ Generate Full SOW Document"):
        if not reference_text:
            st.warning("⚠️ Please upload an input document.")
        else:
            buffer = generate_gts_sow(client, model_name, reference_text,client_name)

            st.success(f"✅ Full GTS RFP/SOW document generated.")
            st.download_button(
                label="📥 Download Generated Document (.docx)",
                data=buffer,
                file_name=f"GTS_SOW.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
