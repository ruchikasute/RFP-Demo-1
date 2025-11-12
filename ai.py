import streamlit as st
from openai import AzureOpenAI
import os, io
from docx import Document
from PyPDF2 import PdfReader
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import re



def insert_formatted_text(doc, placeholder, raw_text):
    """
    Smart parser: handles any LLM output format (Markdown, numbered, or plain text)
    Converts to clean formatted Word content (Headings, bullets, tables, paragraphs).
    """

    def set_cell_shading(cell, fill_color):
        """Add shading to a table cell."""
        tc_pr = cell._element.tcPr
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)
        tc_pr.append(shd)

    def insert_styled_table(parent, headers, rows):
        """Insert styled table with blue header."""
        table = parent.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h.strip()
            set_cell_shading(hdr_cells[i], "008FD3")
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        for r, row_data in enumerate(rows):
            cells = table.rows[r + 1].cells
            for c, val in enumerate(row_data):
                cells[c].text = str(val).strip()
                set_cell_shading(cells[c], "E7EEF7")
        return table

    def add_heading(doc, text, level):
        """Add heading with correct style and spacing."""
        p = doc.add_paragraph(text.strip(), style=f"Heading {level}")
        p.paragraph_format.space_after = Pt(4)
        return p

    # --- Find placeholder ---
    inserted = False
    for para in doc.paragraphs:
        if placeholder in para.text:
            inserted = True
            parent = para._element.getparent()
            idx = parent.index(para._element)
            parent.remove(para._element)

            lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
            new_elements = []
            i = 0

            while i < len(lines):
                line = lines[i]

                # --- Markdown Headings ---
                if re.match(r"^#{1,3}\s", line):
                    level = line.count("#")
                    heading_text = re.sub(r"^#+\s*", "", line)
                    new_elements.append(add_heading(doc, heading_text, level)._element)
                    i += 1
                    continue

                # --- Numbered headings (1., 2., etc.) ---
                if re.match(r"^\d+\.\s+[A-Z]", line):
                    heading_text = re.sub(r"^\d+\.\s*", "", line)
                    new_elements.append(add_heading(doc, heading_text, 1)._element)
                    i += 1
                    continue

                # --- Bold-only headings ---
                if re.match(r"^\*\*[^\*]+\*\*$", line):
                    heading_text = re.sub(r"\*\*", "", line)
                    new_elements.append(add_heading(doc, heading_text, 2)._element)
                    i += 1
                    continue

                # --- Tables ---
                if line.startswith("|") and "|" in line:
                    table_lines = []
                    while i < len(lines) and lines[i].startswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    headers = [h.strip("* ") for h in table_lines[0].strip("|").split("|")]
                    rows = [r.strip("|").split("|") for r in table_lines[2:]]
                    table = insert_styled_table(doc, headers, rows)
                    new_elements.append(table._element)
                    continue

                # --- Bullets (dash or dot) ---
                if line.startswith("- ") or line.startswith("• "):
                    bullet_text = re.sub(r"^[-•]\s*", "", line)
                    p = doc.add_paragraph(bullet_text, style="List Bullet 2")
                    p.paragraph_format.left_indent = Pt(18)
                    p.paragraph_format.space_after = Pt(2)
                    new_elements.append(p._element)
                    i += 1
                    continue

                # --- Regular text ---
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.2
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                new_elements.append(p._element)
                i += 1

            # insert processed elements
            for el in reversed(new_elements):
                parent.insert(idx, el)
            break

    if not inserted:
        st.warning(f"⚠️ Placeholder {placeholder} not found — appending content at end.")
        doc.add_paragraph(raw_text)




# ============================================================
# Helper: Generate GTS RFP/SOW Response
# ============================================================

def generate_ai_sow(client, model_name):
    """Generate a detailed SOW response document using Azure OpenAI."""
    today = datetime.now().strftime("%d %B %Y")
    template_path = "Template/GTS_Template.docx"

    # --- Load or create Word template ---
    if os.path.exists(template_path):
        doc = Document(template_path)
        st.info("📄 Using Crave Word template.")
    else:
        st.warning("⚠️ Template not found. Creating blank document.")
        doc = Document()
        doc.add_paragraph("<<CONTENT START>>")

    # Pre-clean uploaded RFP text to remove proposal-style headers

    # reference_text = re.sub(r"(?i)(proposal\s+for\s+sap\s+.*?)(\n|$)", "", reference_text)
    # reference_text = re.sub(r"(?i)(^|\n)\s*1\.\s*proposal\s+for\s+.*", "", reference_text)
    # reference_text = re.sub(r"(?i)sap\s+gts\s+processes\s+and\s+enhancements", "", reference_text)

    # --- Build LLM prompt ---
    prompt = f"""
You are a Senior SAP AI consultant from Crave InfoTech preparing a professional
RFP Response / Statement of Work for client.

Now, generate a **comprehensive and polished proposal** structured with the following sections:

1. Introduction
   The Introduction should open the document with a concise, narrative-style overview of the project’s purpose, primary objectives, and key technologies involved.  
   Expand this into 3–4 rich paragraphs that clearly connect the business objectives to Crave InfoTech’s expertise in **SAP AI, Machine Learning, and SAP Business Technology Platform (BTP)**.  
   Describe how Crave leverages its **SAP AI Core and AI Launchpad**, **SAP Build Process Automation (SBPA)**, and **SAP Datasphere** capabilities to deliver actionable intelligence, operational efficiency, and predictive insights.  
   Focus on how Crave InfoTech’s experience in **AI-driven process automation**, **data modeling**, and **intelligent integrations** will enable client to achieve accuracy, scalability, and innovation at enterprise scale.  
   Maintain a formal and consultative tone that positions Crave InfoTech as a trusted partner for intelligent automation and AI adoption within SAP landscapes.

2. Project Scope
   2.1. In-Scope Items  
       - Define the AI implementation scope: model development, integration, training data pipelines, monitoring, and governance.  
   2.2. Solution Architecture  
       - Describe proposed SAP AI architecture, including components like AI Core, AI Launchpad, SAP Datasphere, and SBPA integration.  
   2.3. Prerequisites and Key Assumptions  
       - Outline data readiness, model deployment prerequisites, and environment setup on SAP BTP.  
   2.4. Out of Scope  
       - Clearly list excluded modules, non-SAP integrations, or experimental features.  
   2.5. Project Document Deliverables  
       - Enumerate expected deliverables such as Design Document, Training Dataset Summary, Model Deployment Guide, and User Handbook.

3. Bill of Materials (BOM)  
   - Mention major components, licenses, and tools (e.g., SAP AI Core, SAP Datasphere, TensorFlow, Python SDK) in tabular format.

4. Responsibility Matrix  
   **Part 1 – RACI Table:**  
   Provide a table with the following columns:  
   *Task* | *Crave InfoTech (R/A/C/I)* | *Client (R/A/C/I)*  

   **Part 2 – Roles & Responsibilities Table:**  
   Below the RACI, add a second table titled “Roles & Responsibilities” with:  
   *Role* | *Key Responsibilities*  
   (e.g., AI Architect, Data Engineer, Functional Consultant, Project Manager, QA Lead)

5. Project Delivery Approach  
   - Project Organization Structure  
   - Project Resource Planning  
       Begin with a short paragraph describing how Crave’s Delivery and Project Managers identify and allocate resources for AI lifecycle management—from data engineering and model training to deployment and support.  
       Then generate a detailed table titled **“Exhibit: Project Resource Planning”**, with columns:  
       *Project Phase* | *Functional Consultant* | *Technical Consultant* | *Other Roles (Developer, Tester, PM, Data Scientist, Architect, etc.)*  
       Include realistic project phases:  
       - Project Preparation  
       - Business Blueprint / Detailed Design  
       - Model Development and Training  
       - Validation and Integration Testing  
       - Documentation and User Training  
       - Deployment and Go-Live  
       - Post Go-Live Support  
       Fill each cell with indicative resource involvement.  

   - Implementation Methodology  
       Describe Crave’s methodology combining **SAP Activate** and **AI lifecycle best practices** — data preprocessing, model iteration, explainability, and continuous retraining.  

   - Communication Plan  
       **Part 1 – Communication Schedule Table:**  
       Columns: *Interaction*, *Frequency*, *Purpose* (daily stand-ups, weekly governance calls, monthly reviews).  

       **Part 2 – Issue Management and Escalation Process:**  
       Columns: *Task*, *Timescale*, *Responsibility* (define SLAs and escalation triggers).  

       **Part 3 – Issue Classification Table:**  
       Columns: *Severity Level*, *Definition*, *Reporting Process*, *Solution Responsible* (Low, Serious, Critical).  

       **Part 4 – Escalation Process Table:**  
       Columns: *Issue Type*, *Escalation Point*, *Escalation Criteria* (e.g., unresolved issues, model performance gaps, or missed delivery timelines).  
       Populate realistic governance levels such as:  
       - Governance Role (Project Core Group)  
       - Project Delivery Manager  
       - Crave Technology Project Manager  
       - Crave Technology Delivery Manager  

6. Timelines  
   The overall delivery duration should be around **3 months (12 weeks)**.  
   Provide 1 concise paragraph describing Crave InfoTech’s milestone-based delivery approach ensuring quality and on-time completion.

   6.1. Delivery Timeliness  
        - Mention total duration (≈12 weeks) and Crave’s structured phase-wise governance.

   6.2. Efforts and Resource Allocation  
        - Include a brief **table** with columns: *Phase*, *Duration (Weeks)*, *Key Activities*, *Indicative Resources*.  
        - Example phases: Kickoff, Design, Development, Testing, Go-Live, Post-Go-Live.  
        - Keep descriptions short (1 line each).
 
   6.4. Payment Terms  
       Present typical milestones in tabular format (e.g., Kickoff, Design Sign-off, Model Deployment, UAT Completion, Go-Live).

7. Sign-Off  
   Add formal sign-off language ensuring mutual agreement on scope, deliverables, and timelines.

8. Other Assumptions  
   Highlight additional assumptions for clarity.  
   8.1. Dependency – Specify dependencies like client-provided data, test environments, or API access.  
   8.2. Limitations – Mention limitations of model accuracy, third-party data quality, or change control.  
   8.3. General Provisions – Include standard terms, confidentiality, and governance provisions.  

Ensure:
- Content flows logically and professionally.
- Use **SAP AI**, **SAP BTP**, and **SBPA** terminology appropriately.
- Tone: confident, formal, and consultative.
- Avoid bullet overload; prefer paragraph narrative where possible.
- The document **must begin directly with the section heading “1. Introduction”**, without any preamble or title like “Proposal for …”.

"""
    # --- Call Azure LLM ---
    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4
        )
        sow_text = response.choices[0].message.content.strip()
        # --- 🧹 Clean unwanted headers ---
        cleanup_patterns = [
            r"(?i)^crave\s+infotech\s+proposal.*\n?",
            r"(?i)^sap\s+gts\s+implementation.*\n?",
            r"(?i)^proposal\s+for.*\n?",
            r"(?i)^client\s*:\s*.*\n?",
            r"(?i)^date\s*:\s*.*\n?",
            r"(?i)^---+\s*\n?",
        ]
        for pattern in cleanup_patterns:
            sow_text = re.sub(pattern, "", sow_text).strip()

        # --- Force start from Introduction ---
        intro_match = re.search(r"(?i)(^|\n)(\d+\.\s*)?introduction", sow_text)
        if intro_match:
            sow_text = sow_text[intro_match.start():].strip()
    except Exception as e:
        st.error(f"⚠️ Error calling model: {e}")
        sow_text = "Error generating document content."


        # --- Insert formatted content instead of plain text ---
    insert_formatted_text(doc, "<<CONTENT START>>", sow_text)


    # --- Add client name and metadata at top ---
    doc.add_paragraph()


    # --- Save to memory ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ============================================================
# Streamlit UI
# ============================================================

def main():
    st.title("🌐 AI — SOW Generator")

    # uploaded_file = st.file_uploader(
    #     " ",
    #     type=["pdf", "docx"],
    #     key="rfp_uploader",
    #     help="Upload the client RFP or reference document (PDF/DOCX).",
    #     label_visibility="collapsed"
    # )

    # Azure setup
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )
    model_name = "codetest"

    if st.button("⚡ Generate Full SOW Document"):

        buffer = generate_ai_sow(client, model_name)

        st.success(f"✅ Full AI RFP/SOW document generated.")
        st.download_button(
                label="📥 Download Generated Document (.docx)",
                data=buffer,
                file_name=f"AI_SOW.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
