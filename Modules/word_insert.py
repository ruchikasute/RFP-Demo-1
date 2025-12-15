import re
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os, io
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def remove_static_before_placeholder(doc, placeholder):
    for i, p in enumerate(doc.paragraphs):
        if placeholder in p.text:
            # Delete only the 2–3 paragraphs immediately above the placeholder
            start = max(0, i - 3)
            for j in range(i - 1, start - 1, -1):
                if placeholder not in doc.paragraphs[j].text:
                    p_to_remove = doc.paragraphs[j]._element
                    p_to_remove.getparent().remove(p_to_remove)
            return


def _create_paragraph_after(existing_paragraph, text=None, apply_default_spacing=True):
    new_p_elm = OxmlElement("w:p")
    existing_paragraph._p.addnext(new_p_elm)
    new_para = Paragraph(new_p_elm, existing_paragraph._parent)

    if text:
        run = new_para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)

    pf = new_para.paragraph_format

    if apply_default_spacing:
        pf.space_before = Pt(4)
        pf.space_after = Pt(6)
    else:
        # For headings: let the Word template decide spacing
        pf.space_before = None
        pf.space_after = None

    return new_para


# import re
def style_has_numbering(doc, style_name):
    try:
        style = doc.styles[style_name]   # <-- FIX HERE
    except KeyError:
        return False

    try:
        return "numPr" in style._element.xml
    except:
        return False


def extract_block(tag, text):
    pattern = rf"<{tag}>(.*?)(?=<[A-Z][A-Z_ ]+>|$)"

    match = re.search(pattern, text, flags=re.DOTALL)
    if not match:
        return ""
    return match.group(1).strip()


# def insert_formatted_text(doc, placeholder, text, resource_table_markdown_text=None):
#     """
#     Inserts formatted content into the Word doc based on RAW line patterns.
#     Supports:
#     - Markdown Headings (#, ##, ###)
#     - Bullets (-, •, *)
#     - Nested bullets (2+ spaces)
#     - Markdown tables
#     - Bold markers (**text**)
#     """
def insert_formatted_text(
    doc,
    placeholder,
    text,
    resource_table_markdown_text=None,
    replace=False,
    **kwargs
):

    # ---- FIND PLACEHOLDER ----
    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_p = p
            break

    if not target_p:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            target_p = p
                            break

    if not target_p:
        return

    # target_p.text = ""
    # # last_para = target_p
    # lines = text.splitlines()
    # # Replace placeholder with empty text but KEEP paragraph
    # target_p.text = target_p.text.replace(placeholder, "")
    # last_para = target_p

    # --------------------------------------------------------
    # NEW PATCH — CLEAR ORIGINAL TEMPLATE TEXT
    # --------------------------------------------------------
    # if replace:
    #     target_p.text = ""   # wipe old static content completely

    if replace:
        # 1️⃣ Completely remove the entire paragraph node (not just text)
        p = target_p._p                   # XML paragraph <w:p>
        parent = p.getparent()
        idx = parent.index(p)
        parent.remove(p)

        # 2️⃣ Create a brand-new empty paragraph at the same index
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        parent.insert(idx, new_p)
        target_p = Paragraph(new_p, doc)

    
    # last_para remembers where we add new paragraphs
    last_para = target_p

    lines = text.splitlines()

    # Detect whether List Bullet styles have numbering
    list_bullet_has_num = style_has_numbering(doc, "List Bullet")
    list_bullet2_has_num = style_has_numbering(doc, "List Bullet 2")



    i = 0

    # ----------------------------------------
    # MAIN LOOP
    # ----------------------------------------
    while i < len(lines):

        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        # ----------------------------------------------------
        # 1. DETECT MARKDOWN TABLE
        # ----------------------------------------------------
        if "|" in stripped and stripped.count("|") >= 2:
            table_lines = [stripped]
            j = i + 1

            while j < len(lines) and "|" in lines[j]:
                table_lines.append(lines[j].strip())
                j += 1

            insert_markdown_table_after(doc, last_para, "\n".join(table_lines))
            i = j
            continue

        # ----------------------------------------
        # 2. DETECT MARKDOWN HEADINGS
        # ----------------------------------------
        is_md_h1 = stripped.startswith("# ")
        is_md_h2 = stripped.startswith("## ")
        is_md_h3 = stripped.startswith("### ")
        is_md_h4 = stripped.startswith("#### ") 

        if is_md_h4:
            clean = stripped.replace("#### ", "").strip()
            para = _create_paragraph_after(last_para, clean)

            # Apply your Heading Style for these subheadings
            para.style = "Heading 4"

            last_para = para
            i += 1
            continue

        
        if is_md_h3:
            clean = stripped.replace("### ", "").strip()
            para = _create_paragraph_after(last_para, clean)
            
            para.style = "Heading 3"
            last_para = para
            i += 1
            continue

        if is_md_h2:
            clean = stripped.replace("## ", "").strip()
            para = _create_paragraph_after(last_para, clean)


            para.style = "Heading 2"
            last_para = para
            i += 1
            continue

        if is_md_h1:
            clean = stripped.replace("# ", "").strip()
            para = _create_paragraph_after(last_para, clean)

            para.style = "Heading 1"
            last_para = para
            i += 1
            continue

        # --------------------------------------------------------
        # FORCE PARAGRAPH BREAK AFTER ANY LINE ENDING WITH ":"
        # --------------------------------------------------------
        if stripped.endswith(":"):
            # insert this line as its own paragraph
            para = _create_paragraph_after(last_para, stripped)
            para.style = "Normal"
            last_para = para

            # insert a blank line to separate bullets
            last_para = _create_paragraph_after(last_para, "")
            i += 1
            continue

        # ----------------------------------------
        # DETECT NUMBERED HEADINGS (H2 and H3)
        # ----------------------------------------
        num_heading = re.match(r"^(\d+(?:\.\d+)+)\s+(.+)$", stripped)

        if num_heading:
            numbering = num_heading.group(1)
            title_text = num_heading.group(2)

            level = numbering.count(".") + 1

            # Heading paragraph — DO NOT apply python spacing
            para = _create_paragraph_after(
                last_para, stripped, apply_default_spacing=False
            )

            if level == 2:
                para.style = "Style2"
            else:
                para.style = "Style3"

            last_para = para
            i += 1
            continue



        # ----------------------------------------
        # 3. DETECT BULLETS
        # ----------------------------------------
        bullet_match = re.match(r"^\s*([-•*])\s*(.+)$", raw)
        is_bullet = bullet_match is not None

        indent_level = len(raw) - len(raw.lstrip(" "))

        # ----------------------------------------
        # 4. CREATE PARAGRAPH WITH CORRECT STYLE
        # ----------------------------------------
        if is_bullet:

            # Insert blank when previous is not bullet
            if i > 0:
                prev = lines[i - 1].strip()
                # if prev == "" or not re.match(r"^\s*[-•*]\s*", prev):
                #     last_para = _create_paragraph_after(last_para, "")

            # Create new bullet paragraph
            para = _create_paragraph_after(last_para, "")
            
            # # APPLY REAL WORD BULLET STYLE
            if indent_level < 2:
                para.style = "List Bullet 2"
            else:
                para.style = "List Bullet 2"


            # DO NOT REMOVE NUMBERING RUNS – Word needs them
            for r in list(para._p.iter("w:r")):
                txt = "".join(t.text for t in r.iter("w:t"))
                if txt.strip() != "":
                    para._p.remove(r)
            # This preserves <w:numPr> so Word bullets appear correctly

            # Extract the actual bullet text
            _, clean = bullet_match.groups()
            clean = clean.strip()

            # Bold or normal text
            if "**" in clean:
                segments = re.split(r"\*\*(.+?)\*\*", clean)
                for idx, seg in enumerate(segments):
                    run = para.add_run(seg)
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    if idx % 2 == 1:
                        run.bold = True
            else:
                run = para.add_run(clean)
                run.font.name = "Arial"
                run.font.size = Pt(11)

        else:
            # Normal paragraph
            para = _create_paragraph_after(last_para, stripped)
            para.style = "Normal"

        last_para = para
        i += 1
        


def insert_markdown_table_after(doc, last_para, table_text):
    """
    Converts a markdown-style table into a Word table and inserts it after last_para.
    Entire formatting comes from Word style 'Style1'.
    """

    rows = [r.strip() for r in table_text.split("\n") if r.strip()]
    if len(rows) < 2:
        return None

    # Parse header + data
    headers = [h.strip() for h in rows[0].strip("| ").split("|")]
    data_rows = [
        [c.strip() for c in r.strip("| ").split("|")]
        for r in rows[2:]
    ]

    # Create table (Word will apply Style1 styling)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Style1"   # 🔥 Apply your custom table style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- HEADER ---
    hdr_cells = table.rows[0].cells
    for ci, head in enumerate(headers):
        hdr_cells[ci].text = head
        # Make header bold (Style1 may already do this)
        for run in hdr_cells[ci].paragraphs[0].runs:
            run.font.bold = True
        hdr_cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # --- BODY ROWS ---
    for r in data_rows:
        row_cells = table.add_row().cells
        for ci, cell_text in enumerate(r):
            if ci < len(row_cells):
                row_cells[ci].text = cell_text

    # Insert table in the document after placeholder paragraph
    last_para._p.addnext(table._tbl)

    return table

from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def insert_image_at_placeholder(doc, placeholder, image_bytes):
    """
    Finds <PPT_IMAGE> placeholder and replaces it with an image.
    """
    target_para = None

    # 1. Search in normal paragraphs
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_para = p
            break

    # 2. Search in tables if not in body
    if not target_para:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            target_para = p
                            break

    if not target_para:
        return  # placeholder not found

    # Delete placeholder text
    target_para.text = ""

    # Insert image in the same paragraph
    run = target_para.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(7))   # adjust size

    # Optional styling
    target_para.alignment = 1   # center alignment


def insert_plain_preview(doc, placeholder, preview_text):
    """
    Inserts text exactly as shown in preview:
    - empty lines preserved
    - headings detected as plain bold text (if surrounded by **)
    - paragraphs inserted as-is
    """
    target_p = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_p = p
            break

    if not target_p:
        return

    # clear placeholder paragraph
    target_p.text = ""
    last = target_p

    blocks = preview_text.split("\n\n")

    for block in blocks:
        lines = block.strip().split("\n")

        # If block starts with **Heading**
        if len(lines) == 1 and lines[0].startswith("**") and lines[0].endswith("**"):
            heading = lines[0][2:-2].strip()
            para = _create_paragraph_after(last, heading)
            para.style = "Heading 3"
            last = para
            continue

        # Normal paragraph
        para = _create_paragraph_after(last, block.strip())
        para.style = "Normal"
        last = para


def normalize_heading_indents(doc):
    """
    Fixes huge indentation by resetting Heading 2 and Heading 3 paragraph indents.
    """
    for heading_style in ["Heading 2", "Heading 3"]:
        style = doc.styles[heading_style]
        pf = style.paragraph_format
        
        pf.left_indent = Pt(0)            # No left indent
        pf.first_line_indent = Pt(0)      # No first-line indent
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        # YOU CAN TWEAK THIS FOR YOUR TEMPLATE

from docx.shared import Pt, RGBColor

def update_heading2_style(doc):
    """
    Overrides the Heading 2 style formatting while keeping it as Heading 2.
    """

    style = doc.styles["Heading 2"]
    font = style.font

    # Customize here:
    font.name = "Arial"
    font.size = Pt(14)                 # Font size
    font.bold = True                   # Optional
    font.color.rgb = RGBColor(0, 32, 96)   # Dark blue numbers, adjust as needed

    # Adjust indentation
    p_format = style.paragraph_format
    p_format.left_indent = Pt(0)       # remove huge indent
    p_format.first_line_indent = Pt(0)
    p_format.space_before = Pt(15)
    p_format.space_after = Pt(4)


from docx.shared import Pt
from docx.oxml import OxmlElement

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def insert_table_at_placeholder(doc, placeholder, df):
    """
    Reliable table insertion that:
    - Finds placeholder text (even inside bullets or merged paragraphs)
    - Removes ONLY the placeholder text, not the paragraph
    - Inserts a clean anchor paragraph
    - Inserts the table after the anchor
    - Prevents Word XML merging issues
    """

    target_para = None

    # 1. Search paragraphs normally
    for p in doc.paragraphs:
        if placeholder in p.text:
            target_para = p
            break

    # 2. Search inside tables
    if not target_para:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            target_para = p
                            break

    if not target_para:
        return False  # placeholder not found

    # 3. Remove ONLY the placeholder text from the paragraph
    target_para.text = target_para.text.replace(placeholder, "").strip()

    # 4. Create a NEW anchor paragraph just after the placeholder paragraph
    anchor_xml = OxmlElement("w:p")
    target_para._p.addnext(anchor_xml)
    anchor_para = Paragraph(anchor_xml, doc)

    # 5. Now build the table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Style1"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    # Body rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # 6. Insert the table after anchor paragraph
    anchor_xml.addnext(table._tbl)

    return True


# def insert_table_at_placeholder(doc, placeholder, df):
#     """
#     Inserts a Word table exactly where the placeholder appears in the text content
#     generated by the LLM (same behavior as insert_image_at_placeholder).
#     """

#     target_para = None

#     # 1. Search paragraphs
#     for p in doc.paragraphs:
#         if placeholder in p.text:
#             target_para = p
#             break

#     # 2. Search inside tables
#     if not target_para:
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for p in cell.paragraphs:
#                         if placeholder in p.text:
#                             target_para = p
#                             break

#     if not target_para:
#         return False  # placeholder missing

#     # 3. Remove placeholder text
#     # target_para.text = ""
#     p = target_para._p
#     parent = p.getparent()
#     idx = parent.index(p)
#     parent.remove(p)

#     # Create a clean anchor paragraph
#     new_p = OxmlElement("w:p")
#     parent.insert(idx, new_p)
#     anchor_para = Paragraph(new_p, doc)

#     # 4. Create table object
#     table = doc.add_table(rows=1, cols=len(df.columns))
#     table.style = "Style1"        # uses your existing table style
#     hdr_cells = table.rows[0].cells

#     # Header row
#     for i, col in enumerate(df.columns):
#         hdr_cells[i].text = str(col)
#         for run in hdr_cells[i].paragraphs[0].runs:
#             run.font.bold = True
#         hdr_cells[i].paragraphs[0].font = Pt(10)

#     # Body rows
#     for _, row in df.iterrows():
#         row_cells = table.add_row().cells
#         for i, value in enumerate(row):
#             row_cells[i].text = str(value)

#     # 5. Insert table XML after placeholder paragraph
#     target_para._p.addnext(table._tbl)

#     return True


def remove_table_by_tag(doc, tag):
    """
    Removes the FIRST table where ANY run in ANY paragraph 
    of ANY cell contains the placeholder tag.
    """

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Check each run inside each paragraph
                for p in cell.paragraphs:
                    full_text = "".join(run.text for run in p.runs)
                    if tag in full_text:
                        # Remove entire table
                        tbl = table._element
                        tbl.getparent().remove(tbl)
                        return True
    return False


# def remove_table_by_tag(doc, tag):
#     for table in doc.tables:
#         try:
#             if tag in table.cell(0, 0).text:
#                 tbl = table._element
#                 tbl.getparent().remove(tbl)
#                 return True
#         except:
#             pass
#     return False
