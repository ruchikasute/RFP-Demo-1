import streamlit as st
from openai import AzureOpenAI
import os, io, re
from datetime import datetime
from docx import Document

# Assuming these imports work as in your original code
# Run init ONLY once
from Modules.docx_table_handler import update_table_in_doc


from Modules.extractors import (
    extract_text_from_file,
    summarize_large_rfp,
    extract_image_from_slide,
    extract_table_from_slide,
    extract_slide7_summary,
    extract_total_interfaces_from_slide,
)

from Modules.word_insert import (
    insert_formatted_text,
    insert_image_at_placeholder,
    insert_table_at_placeholder,
    remove_table_by_tag,
    remove_static_before_placeholder, 
)

from Modules.placeholders import (
    replace_client_name_in_doc,
    replace_submission_date,
    replace_inline_placeholder,
    insert_document_number,
    generate_document_number,
)

from Modules.preview import section_preview_tabs


def load_template_section(section_name):
    template_path = "Template/Integration_TemplateV3.docx"
    doc = Document(template_path)

    section_text = []
    capture = False

    for p in doc.paragraphs:
        line = p.text.strip()

        # Start capturing AFTER the heading
        if line == section_name:
            capture = True
            continue   # <-- FIX: Skip heading itself

        # Stop at next section
        elif capture and line in [
            "About Crave InfoTech",
            "Our Understanding",
            "Project Scope",
            "Solution",
            "Project Approach",
            "Resource Allocation",
            "Timelines",
            "Commercials & Payment Terms",
            "Governance",
            "Sign Off",
            "Key Assumptions"
        ]:
            break

        if capture:
            section_text.append(line)

    return "\n".join(section_text)


def remove_text_at_placeholder(doc, placeholder):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, "")

def remove_static_after_placeholder(doc, placeholder):
    remove = False
    for p in doc.paragraphs:
        if placeholder in p.text:
            remove = True
            continue
        if remove:
            if p.text.strip().startswith("2.") or p.text.strip().startswith("About Crave"):  
                break
            p.text = ""

import re

def clean_rag_text(text):
    # Remove page numbers, headers, "Company Confidential", weird spacing
    text = re.sub(r"Company Confidential.*?\n", "", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)

    # Remove table garbage or repeated tokens
    text = re.sub(r"[•▪►➢●■]+", "-", text)

    # Fix broken words like "Mo erni ation"
    text = re.sub(r"([A-Za-z])\s+([A-Za-z])", r"\1\2", text)

    return text.strip()


def generate_executive_summary(client, model_name, reference_text, client_name, total_interfaces):
    from Modules.knowledge import similarity_search

    # 1. RAG: Extract ONLY success stories
    success_chunks = similarity_search(
        query="Crave success stories SAP Integration Suite migration automation Merck Eli Lilly Meijer case study",
        category="Integration",
        top_k=4
    )

    success_text = "\n".join(success_chunks) if success_chunks else ""

    prompt = f"""
Generate ONLY the dynamic part of the Executive Summary for {client_name}.
Do NOT rewrite or duplicate the static Crave branding paragraphs already in the template.

CLIENT: {client_name}
TOTAL INTERFACES: {total_interfaces}

RFP CONTEXT (background only, do not quote directly):
{reference_text}

REAL SUCCESS STORIES (STRICTLY factual; do NOT invent names/metrics):
{success_text}

WRITE ONLY THESE FOUR SECTIONS (NO HEADINGS):

1. Client Context Paragraph
   - Write a long, detailed paragraph (minimum 120–150 words)
   - Break into 2–3 natural sentences per line when needed


2. Business Value Paragraph
   - Minimum 150–180 words
   - Write in long-form consulting style
   - Cover operational, strategic, performance, governance, scalability benefits

3. Success Stories (short intro + bullets)
   - First: Write ONE sentence introducing Crave’s proven track record
     (do NOT repeat static branding)
   - Each bullet MUST be taken directly from the real success stories above
   - Do NOT expand, reinterpret, or add new facts
   - Do NOT write paragraphs here — ONLY bullets

4. Closing Paragraph
   - Minimum 120 words
   - Why Crave is the right partner
   - Emphasize governance, predictable delivery, and scalability

RULES:
- No markdown
- No headings
- Bullets ARE allowed. Use simple dash (-) or dot (•).
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role":"user", "content": prompt}],
        temperature=0.45
    )

    return response.choices[0].message.content.strip()

def generate_why_crave(client, model_name, reference_text, client_name):
    from Modules.knowledge import similarity_search

    # Success stories RAG

    success_chunks = similarity_search(
        query = (
        "GxP FDA 21 CFR Part 11 data integrity compliance "
        "Life Sciences Success Stories Global Pharmaceutical Leader "
        "API-first event-driven architecture SAP Edge Integration Cell "
        "Advanced Event Mesh regulatory environments "
        "BODS-to-HANA IoT-based medical devices Datasphere"
    )

    ,
        category="Integration",
        top_k=4
        )
    long_chunks = [c for c in success_chunks if len(c) > 500]
    if not long_chunks:
        fallback = similarity_search(
            query=(
                "Life Sciences Success Stories detailed narrative "
                "Eli Lilly Merck migration roadmap long page section "
                "integration modernization text block"
            ),
            category="Integration",
            top_k=8
        )
        long_chunks = [c for c in fallback if len(c) > 500]

    clean_long_chunks = [clean_rag_text(c) for c in long_chunks]
    success_text = "\n".join(clean_long_chunks)

    # # Accelerator RAG
    accelerator_chunks = similarity_search(
        query="Crave Integration Workbench CIW Migration Factory Cloud ALM Testing Automation Framework accelerators differentiators",
        category="Integration",
        top_k=3
    )
    accelerator_text = "\n".join(accelerator_chunks) if accelerator_chunks else ""

    prompt = f"""
Generate a full proposal-style “Why Crave” section for {client_name}.  
Follow the structure and style of the Merck sample provided below.  
Write in a consulting, enterprise-grade tone with rich narrative detail.

==================================================
REAL SUCCESS STORIES (USE ONLY THESE)
==================================================
{success_text}

==================================================
REAL DIFFERENTIATORS & ACCELERATORS
==================================================
{accelerator_text}


STRICT OUTPUT STRUCTURE (DO NOT CHANGE HEADINGS):

## Why Crave 

#### Alignment with {client_name}’s Integration Vision  
Write 130–160 words explaining why Crave aligns with {client_name}'s modernization roadmap, governance needs, hybrid landscape, and SAP Integration Suite adoption.  
Follow the writing style of “Alignment with Merck Group’s Integration Vision”.

#### Crave’s Differentiators – Accelerators and Tools  
Write a short intro paragraph (60–80 words).  
Then provide these exact bullets and its description in the same line:  
• Crave Integration Workbench (CIW): 
• Migration Factory Framework: 
• Cloud ALM Dashboards:  
• Testing Automation Framework:  

#### Crave InfoTech’s Experience in the Industry  
Write 120–150 words describing domain expertise, compliance needs, regulatory frameworks, reusable patterns, and delivery maturity.

#### Deep Domain Expertise  
Write 100–130 words explaining Crave’s deep domain capabilities, validation experience, compliance requirements, audit-readiness, and governance processes.

#### Success Stories  
Write the full RAG success stories as narrative blocks (no bullets).  
For each story output:
TITLE  
Description  
[[IMAGE:IDENTIFIER]]

Do not shorten. Do not use bullets. 

#### Value Delivered  
Write 4–6 bullets using grounded outcomes drawn from the success stories and accelerator facts.

#### Why This Matters for {client_name}  
Write 130–160 words explaining why Crave’s approach reduces risk, accelerates modernization, strengthens governance, increases transparency, and ensures long-term enterprise value.

RULES:   
- Use the headings exactly.  
- Bullets allowed.  
- No invented facts.  
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role":"user", "content": prompt}],
        temperature=0.35
    )

    return response.choices[0].message.content.strip()


def generate_about_crave(client, model_name, reference_text, client_name):
    """Generate ONLY About Crave InfoTech section"""
    prompt = f"""
You are writing the "About Crave InfoTech" section for {client_name}.

REFERENCE STYLE GUIDE:
{st.session_state.get("knowledge_text", "")}

INSTRUCTIONS:
- Describe Crave InfoTech's expertise in SAP Integration in 2-3 lines
- Highlight migration factory experience
- Professional tone showcasing credibility

Output ONLY the "About Crave InfoTech" content. No tags, no extra text.
"""
    
    response = client.chat.completions.create(
        model="Codetest",
        messages=[{"role":"user","content":prompt}],
        temperature=0.4
    )
    return response.choices[0].message.content.strip()


# def generate_our_understanding_solution(client, model_name, reference_text, client_name, total_interfaces):
#     prompt = f"""
# You are a Senior SAP Integration Consultant from Crave InfoTech.

# Generate the full “Our Understanding & Solution” section for {client_name} using the following three subsections:

# 3.1 Current Landscape
# 3.2 Objectives
# 3.3 Solution

# Use the reference_text ONLY to understand the client's domain, challenges, processes, and expectations.

# REFERENCE TEXT:
# {reference_text}

# ===========================
# STRICT CONTENT RULES
# ===========================
# - Each subsection must be exactly one paragraph of 4–6 lines.
# - Mention the total interfaces ({total_interfaces}) only once inside the 3.1 paragraph.
# - Strategic, business-focused tone (consulting language).
# - No technical terminology: no adapters, mappings, queues, APIs, connectivity, or configuration.
# - No repeated ideas across subsections.
# - Challenges must be pure business challenges (NOT technical).

# ===========================
# CONTENT GUIDANCE
# ===========================

# 3.1 Current Landscape
# Paragraph:
# - Describe business landscape, operational complexity, modernization need, and overall integration environment.
# - Mention total interfaces ({total_interfaces}) exactly once.
# - Business tone only, no technical terms.

# Challenges:
# - After the paragraph, include 3–4 business challenges as hyphen (-) bullets.
# - Keep bullets short and outcome-focused.

# 3.2 Objectives
# - Explain strategic motivations: scalability, operational excellence, governance, cloud readiness, improved delivery consistency.
# - High-level business reasoning only.

# 3.3 Solution
# - Describe transformation vision and business value.
# - Focus on streamlined operations, governance, and future readiness.
# - Avoid technical steps completely.

# ===========================
# FORMAT RULES (CRITICAL)
# ===========================
# - The headings below must appear EXACTLY at the start of their lines:

# 3.1 Current Landscape
# 3.2 Objectives
# 3.3 Proposed Solution

# - No spaces, no tabs, and no invisible Unicode characters before the numbers.
# - Only plain ASCII characters.
# - Do NOT indent headings.
# - Do NOT use markdown (**text**, ###, etc.).
# - Output exactly in this sequence.

# Output only the final content.
# """

#     response = client.chat.completions.create(
#         model=model_name,
#         messages=[{"role": "user", "content": prompt}],
#         temperature=0.4
#     )

#     return response.choices[0].message.content.strip()
from Modules.knowledge import similarity_search
def generate_our_understanding_solution(client, model_name, reference_text, client_name, total_interfaces):
    

    # Small RAG retrieval (minimal size to avoid 429)
    understanding_chunks = similarity_search(
        query="enterprise integration modernization business drivers cloud readiness governance scalability",
        category="Integration",
        top_k=2
    )
    understanding_text = " ".join(understanding_chunks)

    prompt = f"""
You are a Senior SAP Integration Consultant.

Write the “Our Understanding & Solution” section for {client_name} with three subsections:
3.1 Current Landscape
3.2 Objectives
3.3 Recommended Approach

Use the reference_text and the short knowledge text below only for background:
RAG: {understanding_text}
Reference: {reference_text}

RULES:
- Each subsection must be 90–120 words, single paragraph.
- Mention total interfaces ({total_interfaces}) exactly once in 3.1.
- Business language only (no technical terms like adapters, APIs, mappings).
- Tone must resemble consulting narrative seen in enterprise proposals.
- No repeated ideas across subsections.
- After 3.1, include 3 short business challenges as hyphen bullets.
- Do not add extra sections.

CONTENT GUIDANCE:
3.1 = Describe business landscape, operational complexity, modernization need.
3.2 = Describe strategic motivations (scalability, governance, cloud readiness, consistency).
3.3 = Describe transformation vision (streamlining, governance, future readiness).

OUTPUT:
Write ONLY the three sections with exact headings.
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.35
    )

    return response.choices[0].message.content.strip()



def generate_project_scope(client, model_name, reference_text, client_name, total_interfaces):
    
    rag_chunks = similarity_search(
        query="scope assumptions access SPOC UAT responsibility SAP PI PO to Integration Suite business constraints",
        category="Integration",
        top_k=3
    )

    rag_text = " ".join(rag_chunks)
    placeholder2 = "[[REFER ASSESSMENT TABLE BELOW]]"
    prompt = f"""
You are writing the "Project Scope" section for {client_name}'s SAP Integration Suite migration.

CRITICAL:
- This is a Statement of Work (SOW). Keep the content HIGH-LEVEL.
- DO NOT include technical details like file adapters, mappings, Java code, credential migration, etc.
- DO NOT mention complexity categories (Migrate/Adapt/Evaluate).
- DO NOT include numbers except the total interface count {total_interfaces}.
- DO NOT add any paragraphs outside 4.1, 4.2, 4.3, 4.4.
- NO bold (**text**), NO markdown, NO numbering styles.

Write ONLY these four sub-sections in order:

4.1 Migrating SAP PIPO Integration
- One single high-level paragraph (90-130 words, business tone).
- State that migration covers SAP PI/PO to SAP Integration Suite.
- Mention the total interfaces count: {total_interfaces}.
- After the paragraph, output this placeholder on a new line:
{placeholder2}

Then immediately write an "####Assumptions:" followed by 5–7 business-level SOW assumptions in bullet points.

4.2 Deliverables
- All project deliverables
- Documentation, configuration artifacts, testing evidence, KT session

4.3 Acceptance Criteria
- Clear, verifiable criteria for successful migration

4.4 Out of Scope
- Explicit list of exclusions

RULES:
- 4.1 must be a single high-level paragraph (NOT bullets)
- 4.2, 4.3, 4.4 must use bullet points
- Keep bullets short and business-friendly.
- DO NOT invent new sections.
- DO NOT write assessment-style technical details.
- Output ONLY the final content for 4.1 to 4.4.

BEGIN.
"""

    response = client.chat.completions.create(
        model="Codetest",
        messages=[{"role":"user","content":prompt}],
        temperature=0.3
    )

    return response.choices[0].message.content.strip()

def generate_solution_section(client, model_name, reference_text, client_name, total_interfaces):

    placeholder = "[[ARCHITECTURE_IMG]]"

    prompt = f"""
You are creating Section 5 — Solution for the SOW for {client_name}.

MANDATORY:
You MUST include the placeholder EXACTLY as: {placeholder}
Do NOT change it or it becomes invalid.

STRUCTURE TO FOLLOW:

5.1 Proposed Architecture
Write a 6–8 line paragraph describing:
- High-level SAP BTP architecture
- Integration Suite, API-led strategy
- Security & monitoring
- Refer to an architecture diagram

Then on a new line write ONLY:
{placeholder}

5.2 Bill of Material (BOM)
Write 4–6 bullets using (•):
•	SAP BTP Cloud Integration 
•	SAP BTP Cloud Foundry Run Time 
•	SAP BTP Launchpad 
•	SAP BTP Identity Authentication


RULES:
- No markdown
- No bold
- Only paragraphs + bullets (•)
- Output MUST include placeholder {placeholder} exactly once
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
    )

    return response.choices[0].message.content.strip()


import concurrent.futures
def generate_delivery_approach(client, model_name, reference_text, client_name, total_interfaces):

    placeholder = "[[APPROACH_IMG]]"
    knowledge_text = st.session_state.get("knowledge_text", "")
    

    prompt = f"""
You are generating Section 6.1 — Proposed Migration Approach Under Migration Factory Program
for {client_name}'s SAP PI/PO → SAP Integration Suite migration.

==========================
USE BOTH REFERENCES BELOW
==========================

CRAVE KNOWLEDGE REPOSITORY (style, tone, structure):
{knowledge_text}

CLIENT RFP CONTEXT (domain-specific insights, numbers, processes):
{reference_text}

==========================
STRICT STRUCTURE REQUIRED
==========================

6.1 Proposed Migration Approach Under Migration Factory Program

{placeholder}

Write a 2–3 line introduction summarizing the overall migration philosophy.
• Must sound professional, enterprise-grade, Crave’s tone.
• Must reflect context from the RFP (domain, challenges, goals).

Then write EXACTLY 5 MAIN BULLETS (•):
• Discovery
• Assessment
• Migration
• Validation
• Go-Live & Support


Immediately AFTER the “Assessment” bullet and BEFORE writing any sub-bullets,
output this placeholder on a NEW LINE:


Under EACH main bullet, write 2–3 SUB-BULLETS (–):
• Sub-bullets MUST be descriptive (2–3 lines each)
• Sub-bullets MUST reflect real context from client RFP
• No generic placeholder text
• Add meaningful detail (processes, expected outcomes, constraints)
• Sunil likes detailed sub-steps → make them rich & specific

==========================
RULES
==========================
- No markdown (no **, no #, no tables)
- Use only paragraphs, bullets (•), and sub-bullets (–)
- Do NOT add extra sections or numbering
- {placeholder} must appear exactly once
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.45,   # controlled creativity
    )

    return response.choices[0].message.content.strip()

def generate_resource_allocation(client, model_name, reference_text, client_name, total_interfaces):

    placeholder1 = "[[ORGANIZATION_IMG]]"
    placeholder = "[[Refer Activity Table Below]]"
    placeholder3 = "[[Refer Resource Table Below]]"
    placeholder4 = "[[Refer RACI Table Below]]"

    prompt = f"""
You are a Senior SAP Integration Consultant from Crave InfoTech.

Generate ONLY the content for:

7. Resource Allocation
7.1 Project Organization Structure

Client: {client_name}

REFERENCE CONTEXT (for understanding, do not repeat):
{reference_text}

STRUCTURE TO FOLLOW EXACTLY:

7.1 Project Organization Structure
Write a single paragraph (4–6 lines) describing:
- Engagement team structure
- Onsite + offshore working model
- Single point of accountability
- Collaboration between Crave InfoTech and {client_name}
- Governance and quality focus

Then on a NEW LINE output ONLY the placeholder:
{placeholder1}

Paragraph 2:
- Explain how the structure ensures governance, quality, alignment, and coordinated execution
- Highlight communication flows, oversight, and support model
- No operational or technical details

Then on a NEW LINE output ONLY the placeholder:
{placeholder}

7.2 Resource Allocation

Then on a NEW LINE output ONLY the placeholder:
{placeholder3}

Write one paragraph (4–6 lines) describing:
- Crave’s resourcing approach for this project
- How staffing is aligned with phases and delivery timelines
- That Crave’s team deployment summary and {client_name} team responsibilities are provided in the tables below
- Do NOT rewrite or describe the tables
- Business-focused, high-level, no technical details



Then on a NEW LINE output ONLY the placeholder:
{placeholder4}

STRICT RULES:
- No bullets
- No numbering other than "7.1 Project Organization Structure"
- Do NOT mention total interface count
- Do NOT rewrite tables (they already exist in the template)
- Do NOT describe RACI or roles; only the organization structure overview
- Do NOT add text before or after the required output

Output ONLY the required content exactly in this format.
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )

    return response.choices[0].message.content.strip()


def generate_timelines(client, model_name, reference_text, client_name, total_interfaces):

    prompt = f"""
You are a Senior SAP Integration Delivery Manager from Crave InfoTech.

Using the reference text ONLY as background knowledge, generate a refined and business-focused timeline section for {client_name}'s SAP PI/PO to SAP Integration Suite migration. DO NOT quote or copy the reference text.

STRUCTURE THE OUTPUT EXACTLY LIKE THIS:

1. Intro Paragraph (5–7 lines)
   - Explain why phased delivery is required.
   - Explain the logic of sequencing: cleanup → migration → validation → go-live.
   - Mention {total_interfaces} interfaces once.
   - Business tone only, no technical terminology.

2. Phase-Based Timeline (5 paragraphs total — one for each step)
   For each step:
   - 1 paragraph of 4–6 lines.
   - Describe the purpose of the step.
   - Describe estimated duration (use realistic ranges).
   - Explain how it contributes to controlled delivery.
   - No lists, no bullets.

   Suggested durations:
     Discovery: 2–3 weeks  
     Assessment: 3–4 weeks  
     Migration: 10–14 weeks (wave-based execution)  
     Validation: 4–6 weeks  
     Go-Live & Hypercare: 2–4 weeks  

3. Project Plan Image Placeholder
   Add this EXACT line:
   [[PROJECT_PLAN_IMG]]

4. Final Paragraph (4–6 lines)
   - Summarize total project duration (in weeks or months—based on typical 200–300 interface migrations).
   - Explain why wave execution reduces risk.
   - Highlight predictable delivery, governance, and quality assurance.
   - Business tone only.

REFERENCE TEXT (for understanding only, do not quote or copy):
{reference_text}

Rules:
- No markdown (**text**, ###, etc.).
- No lists or bullets.
- No technical words such as adapters, mappings, queues, APIs, connectivity.
- No repeated ideas.
- Write in polished consulting language.
"""
    response = client.chat.completions.create(
        model="Codetest",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4
    )

    return response.choices[0].message.content.strip()

def generate_commercials(client, model_name, client_name):
    placeholder = "[[Refer COMMERCIAL TABLE BELOW]]"
    placeholder2 = "[[Refer PAYMENT TERMS TABLE BELOW]]"

    prompt = f"""
You are generating ONLY the text for Section 9 — Commercials & Payment Terms for {client_name}'s SOW.

IMPORTANT HARD RULES:
- Do NOT output markdown (no ###, no **, no _).
- Do NOT restyle or rename headings.
- Output the placeholders EXACTLY as {placeholder} and {placeholder2}.
- Do NOT move or skip placeholders.
- No tables — the tables exist in the template.
- Tone must be formal and contractual.

===========================
OUTPUT FORMAT (STRICT)
===========================

9.1 Commercials
Paragraph 1: Write a 2–3 line introduction about commercials, pricing structure, and financial terms. Do NOT mention interface counts or specific costs.

On a NEW LINE output ONLY:
{placeholder}

Terms / Notes:
• Pricing includes inactive/redundant ICO cleanup, jointly performed by Crave InfoTech and {client_name}.
• Pricing assumes collaborative effort estimation and requires {client_name} to provide two qualified integration developers throughout the project.
• Prices are exclusive of VAT and any applicable local taxes.
• Project execution is remote; any on-site travel requested by {client_name} will be charged separately on actuals as per FTA norms.
• Delays in approvals, UAT cycles, or dependent activities on {client_name}'s side may impact overall project timelines.

9.2 Payment Terms

On a NEW LINE output ONLY:
{placeholder2}

DO NOT add anything before or after this section.
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )

    return response.choices[0].message.content.strip()

def generate_governance(client, model_name, client_name):
    placeholder1 = "{{REFER TABLE_Interaction BELOW}}"
    placeholder2 = "{{REFER TABLE_Management BELOW}}"
    placeholder3 = "{{REFER TABLE_Classification BELOW}}"
    placeholder4 = "{{REFER TABLE_Process BELOW}}"


    prompt = f"""
You are writing Section 10 — Governance for a Statement of Work (SOW) for {client_name}.

IMPORTANT:
- This is contractual SOW language.
- Maintain formal, professional consulting tone.
- Do NOT remove or rename any section numbers (10.1 to 10.5).
- Do NOT remove or modify any table placeholders.
- Do NOT invent new governance structures.
- Output plain text only (no markdown, no headings symbols).

========================
STRUCTURE TO FOLLOW EXACTLY
========================

10.1 Communication Plan
Write 2–3 professional lines introducing the communication and reporting approach.

Then output this placeholder on a NEW LINE exactly as shown:
{placeholder1}

After the table placeholder, write 1–2 lines explaining that additional meetings may be scheduled as required during the project lifecycle.

10.2 Issue Resolution and Escalation Procedure
Write a clear paragraph explaining:
- Issue identification
- Tracking and resolution
- Escalation ownership
- Crave InfoTech project governance responsibility

Then write this line exactly:
"The responsibilities and timescales for executing these procedures are outlined below:"

On a NEW LINE, output this placeholder exactly:
{placeholder2}

Then write:
"Following guidelines will be followed for reporting and managing issues:"

Provide 6–8 concise bullet points covering:
- Issue reporting
- Logging
- Classification
- Ownership
- Escalation
- Closure

Then write:
"Following exhibit explains the process of issue classification:"

On a NEW LINE, output this placeholder exactly:
{placeholder3}

Then write:
"Following escalation process will be followed as part of the issue management:"

On a NEW LINE, output this placeholder exactly:
{placeholder4}

End this subsection with this sentence:
"The Project Workgroup will comprise representatives from Crave InfoTech and {client_name} key project stakeholders."

10.3 SLA Reporting
Write one structured paragraph describing:
- SLA and KPI monitoring
- Governance cadence (weekly, bi-weekly, monthly)
- Dashboard visibility
- Root cause analysis and corrective actions

10.4 Service Credit Review
Write one paragraph explaining:
- Monthly and quarterly review cycles
- Validation of service credits
- Reporting of deviations
- Corrective action tracking

10.5 Continuous Improvement
Write one paragraph explaining:
- Governance-driven continuous improvement
- Trend analysis
- Preventive actions
- Sustained delivery excellence

========================
STRICT RULES
========================
- Keep numbering exactly as specified.
- Keep placeholders exactly as provided.
- No markdown.
- No tables written manually.
- No extra sections.
- No explanations.

Output ONLY Section 10 content.
"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.25
    )

    return response.choices[0].message.content.strip()


def generate_sign_off(client, model_name, client_name):
    """Generate ONLY Sign Off section"""
    prompt = f"""
You are writing the "Sign Off" section for {client_name}.

INSTRUCTIONS:
- Write formal agreement statement
- Mention both parties (Crave InfoTech and {client_name})
- Include standard legal language for SOW acceptance
- Keep it brief (2-3 sentences)
- Professional and formal tone

Output ONLY the "Sign Off" content. No tags, no extra text.
"""
    
    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role":"user","content":prompt}],
        temperature=0.4
    )
    return response.choices[0].message.content.strip()

def generate_key_assumptions(client, model_name, reference_text, client_name):
    prompt = f"""
You are writing the "Key Assumptions" section for the SAP PI/PO → SAP Integration Suite Migration SOW for {client_name}.

USE THIS CAREFULLY:
REFERENCE TEXT (use it to infer realistic assumptions, constraints, dependencies, and project conditions — do NOT quote it directly):
{reference_text}

STYLE GUIDANCE (use for tone only):
{st.session_state.get("knowledge_text", "")}

REQUIREMENTS:
- Create a complete “Key Assumptions” section for a real SOW.
- You MAY create multiple assumption categories (e.g., Technical Assumptions, IT & Infrastructure Assumptions, Project Delivery Assumptions, Client Responsibilities, Data & Testing Assumptions, Security & Compliance, Third-Party Dependencies, etc.)
- You are FREE to create category names based on what makes sense for this project.
- Each category must have 2–6 bullets depending on relevance.
- Use clean, professional SOW-style language.
- Use project context from reference_text creatively: interface volumes, system constraints, delivery model, client-side effort, migration characteristics, risks, dependencies, etc.
- Do NOT repeat content from other sections (Project Scope, Approach, etc.).
- Do NOT include numbers like 7.1 / 7.2 — the template already handles numbering.
- Just output headings + bullets.
- No paragraphs unless a category genuinely needs one.
- Use markdown headings (## Heading) for each category.
- Bullets must start with •

TONE:
- Business & delivery focused
- Clear, contractual assumptions
- Non-technical wording for business assumptions
- Technical assumptions should be high-level (no protocol-level details)

OUTPUT ONLY:
- Category headings
- Bulleted assumptions (•)

"""

    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4
    )

    return response.choices[0].message.content.strip()


# ========================================
# MAIN GENERATION FUNCTION
# ========================================

def generate_selected_sections(client, model_name, reference_text, client_name, selected_sections):
    """Generate only the sections that user selected"""
    
    # total_interfaces = st.session_state.get("total_interfaces", "UNKNOWN")
    total_interfaces = st.session_state.get("total_interfaces") or "Not Provided"

    generated_sections = {}
    
    # Map NEW section names to their generation functions
    section_generators = {
        "Executive Summary": lambda: generate_executive_summary(client, model_name, reference_text, client_name, total_interfaces),
        "About Crave InfoTech": lambda: generate_about_crave(client, model_name, reference_text, client_name),
        "Our Understanding": lambda: generate_our_understanding_solution(client, model_name, reference_text, client_name, total_interfaces),
        "Project Scope": lambda: generate_project_scope(client, model_name, reference_text, client_name, total_interfaces),
        "Solution": lambda: generate_solution_section(client, model_name, reference_text, client_name, total_interfaces),
        "Project Approach": lambda: generate_delivery_approach(client, model_name, reference_text, client_name, total_interfaces),
        "Resource Allocation": lambda: generate_resource_allocation(client, model_name, reference_text, client_name, total_interfaces),
        "Project Timelines": lambda: generate_timelines(client, model_name, reference_text, client_name, total_interfaces),
        "Commercials & Payment Terms": lambda: generate_commercials(client, model_name, client_name),
        "Governance": lambda: generate_governance(client, model_name, client_name),
        "Sign Off": lambda: generate_sign_off(client, model_name, client_name),
        "Key Assumptions": lambda: generate_key_assumptions(client, model_name, reference_text, client_name),
    }
    

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = {
            executor.submit(section_generators[name]): name
            for name in selected_sections
        }
        
        for future in concurrent.futures.as_completed(futures):
            name = futures[future]
            dynamic_text = future.result()

            if name == "Executive Summary":
                # 1. Generate Why Crave
                why_crave = generate_why_crave(client, model_name, reference_text, client_name)

                # 2. Load the full template block for Executive Summary
                full_template_block = load_template_section("Executive Summary")

                # Use partition to avoid exceptions if placeholder missing
                before, sep, after = full_template_block.partition("<EXEC_SUMMARY>")
                static_intro = before.strip()
                # after contains sustainability + maybe <WHY_CRAVE> + trailing template text

                # If <WHY_CRAVE> exists, split it out; otherwise keep 'after' as sustainability
                sust_part, sep2, rest_after = after.partition("<WHY_CRAVE>")
                sustainability = sust_part.strip() if sep2 else after.strip()

                # 3. Build final content in correct order
                parts = []
                if static_intro:
                    parts.append(static_intro)
                if dynamic_text and dynamic_text.strip():
                    parts.append(dynamic_text.strip())
                if sustainability:
                    parts.append(sustainability)
                if why_crave and why_crave.strip():
                    parts.append(why_crave.strip())

                final_content = "\n\n".join(parts)

                # 4. Preview is exactly this combined content (user edits this block)
                preview_text = final_content

                generated_sections[name] = {
                    "preview": preview_text,
                    "final": final_content,
                    "why_crave": why_crave
                }
            else:
                generated_sections[name] = {
                    "preview": dynamic_text,
                    "final": dynamic_text
                }


        
    
    return generated_sections

from Modules.startup import init
from Modules.docx_table_handler import extract_table_by_tag
# ========================================
# MAIN APP
# ========================================

def main():
    # Title is rendered in the global topbar (main.py). Do not render local st.title to save vertical space.
    st.title("🌐 Integration — SOW Generator")

    if "app_initialized" not in st.session_state:
        
        init("Integration")
        st.session_state["app_initialized"] = True

    st.caption("✨ Restructured sections with selective generation")
    
    # Initialize session state
    st.session_state.setdefault("llm_client", None)
    st.session_state.setdefault("llm_model", None)

    # Azure LLM client (initialize early so generate button inside expander can use it)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )
    model_name = "gpt-4o-mini"

    st.session_state["llm_client"] = client
    st.session_state["llm_model"] = model_name

    # ========================================
    # INPUT CONFIGURATION (COMBINED EXPANDER)
    # ========================================
    with st.expander("⚙️ Input Configuration", expanded=True):
        
        # --- CLIENT DETAILS & RFP UPLOAD ---
        st.markdown("#### 📝 Client Details & RFP Upload")
        
        # Client name input
        client_name = st.text_input("Enter Client Name (required)", "")

        # File upload
        uploaded_file = st.file_uploader(
            "Upload RFP Document",
            type=["pdf", "docx", "xlsx", "pptx"],
            key="rfp_uploader",
            help="Upload PDF, Word, Excel or PowerPoint reference document.",
        )
        
        
        st.markdown("---")
        
        # --- SECTION SELECTION ---
        st.markdown("#### 📋 Select Sections to Generate (in order)")
        st.caption("✨ Tick sections in the order you want them generated. They will be numbered #1, #2, etc.")

        # Master list of sections
        SECTION_LIST = [
            "Executive Summary",
            "About Crave InfoTech",
            "Our Understanding",
            "Project Scope",
            "Solution",
            "Project Approach",
            "Resource Allocation",
            "Project Timelines",
            "Commercials & Payment Terms",
            "Governance",
            "Sign Off",
            "Key Assumptions",
        ]
        
        # Initialize checkbox states in session state (once)
        if "checkbox_states_integration" not in st.session_state:
            st.session_state["checkbox_states_integration"] = {section: False for section in SECTION_LIST}
        
        # Display checkboxes in 2 columns with live order tracking
        col1, col2 = st.columns(2)
        
        # Render checkboxes and track state changes
        for i, section in enumerate(SECTION_LIST):
            col = col1 if i < 6 else col2
            with col:
                # Get current state
                current_state = st.session_state["checkbox_states_integration"][section]
                new_state = st.checkbox(section, value=current_state, key=f"chk_{section}")
                
                # Update state if changed
                st.session_state["checkbox_states_integration"][section] = new_state
        
        # Build selected sections list in the order they appear in SECTION_LIST
        # This preserves the user's selection order
        selected_sections = []
        for section in SECTION_LIST:
            if st.session_state["checkbox_states_integration"][section]:
                selected_sections.append(section)
        
        # Display selected sections with order numbers
        if selected_sections:
            st.markdown("---")
            st.markdown("### ✅ Selected Sections (in generation order)")
            cols_display = st.columns(min(3, len(selected_sections)))
            for idx, section in enumerate(selected_sections):
                with cols_display[idx % len(cols_display)]:
                    st.markdown(f"**#{idx + 1}** — {section}")
        
        st.markdown("---")
        
        # ========================================
        # GENERATE BUTTON (INSIDE EXPANDER)
        # ========================================
        
        if st.button("⚡ Generate Content"):
            st.session_state.pop("edited_sections", None)
            
            # Reset old editor text areas
            for key in list(st.session_state.keys()):
                if key.startswith("editor_"):
                    st.session_state.pop(key)

            # Allow generation even without RFP
            reference_text = st.session_state.get("reference_text", "")
            if not reference_text:
                reference_text = ""        # Use empty reference
                st.info("ℹ No RFP uploaded — generating a generic SOW draft.")

            
            from Modules.llm_executor import submit_llm_job, wait_for_job

            if not selected_sections:
                st.warning("⚠ Please select at least one section to generate.")
            else:
                st.session_state.pop("edited_sections", None)

                generated_sections = {}
                status = st.status("⏳ Generating selected sections...", expanded=True)

                # total interfaces already resolved earlier
                total_interfaces = st.session_state.get("total_interfaces") or "Not Provided"

                # SAME generator map you already have
                section_generators = {
                    "Executive Summary": lambda: generate_executive_summary(
                        client, model_name, reference_text, client_name, total_interfaces
                    ),
                    "About Crave InfoTech": lambda: generate_about_crave(
                        client, model_name, reference_text, client_name
                    ),
                    "Our Understanding": lambda: generate_our_understanding_solution(
                        client, model_name, reference_text, client_name, total_interfaces
                    ),
                    "Project Scope": lambda: generate_project_scope(
                        client, model_name, reference_text, client_name, total_interfaces
                    ),
                    "Solution": lambda: generate_solution_section(
                        client, model_name, reference_text, client_name, total_interfaces
                    ),
                    "Project Approach": lambda: generate_delivery_approach(
                        client, model_name, reference_text, client_name, total_interfaces
                    ),
                    "Resource Allocation": lambda: generate_resource_allocation(
                        client, model_name, reference_text, client_name, total_interfaces
                    ),
                    "Project Timelines": lambda: generate_timelines(
                        client, model_name, reference_text, client_name, total_interfaces
                    ),
                    "Commercials & Payment Terms": lambda: generate_commercials(
                        client, model_name, client_name
                    ),
                    
                    # ✅ ADD THIS — governance treated like every other section
                    "Governance": lambda: generate_governance(client, model_name, client_name),


                    "Sign Off": lambda: generate_sign_off(
                        client, model_name, client_name
                    ),
                    "Key Assumptions": lambda: generate_key_assumptions(
                        client, model_name, reference_text, client_name
                    ),
                }

                jobs = []

                for section_name in selected_sections:

                    # status.write(f"🔄 Queued **{section_name}**")

                    jobs.append(
                        (section_name, submit_llm_job(section_generators[section_name]))
                    )

                    # 🚀 Run 2 at a time
                    if len(jobs) == 2:
                        for sec, job in jobs:
                            result = wait_for_job(job)
                            if sec == "Executive Summary":
                                why_crave = generate_why_crave(client, model_name, reference_text, client_name)

                                full_template_block = load_template_section("Executive Summary")
                                before, _, after = full_template_block.partition("<EXEC_SUMMARY>")
                                static_intro = before.strip()
                                sust_part, sep2, _ = after.partition("<WHY_CRAVE>")
                                sustainability = sust_part.strip() if sep2 else after.strip()

                                parts = []
                                if static_intro:
                                    parts.append(static_intro)
                                if result:
                                    parts.append(result)
                                if sustainability:
                                    parts.append(sustainability)
                                if why_crave:
                                    parts.append(why_crave)

                                final_content = "\n\n".join(parts)

                                generated_sections[sec] = {
                                    "preview": final_content,
                                    "final": final_content,
                                    "why_crave": why_crave
                                }
                            else:
                                generated_sections[sec] = {
                                    "preview": result,
                                    "final": result
                                }


                        #     generated_sections[sec] = {
                        #         "preview": result,
                        #         "final": result
                        #     }
                        jobs.clear()

                # Handle last leftover (if odd count)
                for sec, job in jobs:
                    result = wait_for_job(job)
                    generated_sections[sec] = {
                        "preview": result,
                        "final": result
                    }

                # st.success(f"✅ Generated {len(selected_sections)} sections!")
                # ✅ FINALIZE status ONCE (AFTER loop)
                # import time
                # status.update(
                #     label=f"✅ Generated {len(selected_sections)} sections!",
                #     state="complete"
                # )
                # time.sleep(0.6)
                # status.empty()

                st.success(f"✅ Generated {len(selected_sections)} sections!")

                MASTER_ORDER = [
                    "Executive Summary",
                    "About Crave InfoTech",
                    "Our Understanding",
                    "Project Scope",
                    "Solution", 
                    "Project Approach",
                    "Resource Allocation",
                    "Project Timelines",
                    "Commercials & Payment Terms",
                    "Governance",
                    "Sign Off",
                    "Key Assumptions",
                ]

                ordered_list = []


                for section_name in MASTER_ORDER:
                    if section_name in generated_sections:
                        
                        section_obj = generated_sections[section_name]

                        # ordered_list.append({
                        #     "title": section_name,
                        #     "preview": section_obj["preview"],   # UI
                        #     "final": section_obj["final"]        # DOCX insertion
                        # })
                        entry = {
                            "title": section_name,
                            "preview": section_obj["preview"],
                            "final": section_obj["final"]
                        }

                        # Preserve Why Crave
                        if "why_crave" in section_obj:
                            entry["why_crave"] = section_obj["why_crave"]

                        ordered_list.append(entry)



                st.session_state["edited_sections"] = ordered_list
                # st.success(f"✅ Generated {len(selected_sections)} sections!")

    

    # Extract and process uploaded file
    if uploaded_file and "reference_text" not in st.session_state:
        # Get client_name and uploaded_file from expander context
        # Note: These are already set in the expander above
        raw_text = extract_text_from_file(uploaded_file)
        extracted_items = []
        st.success(f"✅ Extracted {len(raw_text.split())} words")

        # PPT-specific extractions
        if uploaded_file.name.lower().endswith(".pptx"):
            img = extract_image_from_slide(uploaded_file, 17)
            if img:
                image_blob, ext = img
                st.session_state["slide17_image"] = image_blob
                extracted_items.append("📸 Image (Slide 17)")

            table_md = extract_table_from_slide(uploaded_file, 9)
            if table_md:
                st.session_state["slide9_table"] = table_md
                extracted_items.append("📊 Table (Slide 9)")

            table_md_8 = extract_table_from_slide(uploaded_file, 8)
            if table_md_8:
                st.session_state["slide8_table"] = table_md_8
                extracted_items.append("📊 Table (Slide 8)")

            slide7_summary = extract_slide7_summary(uploaded_file)
            if slide7_summary:
                st.session_state["slide7_text"] = slide7_summary
                extracted_items.append("📝 Summary bullets (Slide 7)")

            total = extract_total_interfaces_from_slide(uploaded_file, slide_number=5)
            if total:
                st.session_state["total_interfaces"] = total
                extracted_items.append(f"🔢 Interface count: {total}")

            table_md_18 = extract_table_from_slide(uploaded_file, 18)
            if table_md_18:
                st.session_state["slide18_resources_table"] = table_md_18
                extracted_items.append("👥 Resources table (Slide 18)")

        if extracted_items:
            st.markdown("### 🔎 Extracted assets")
            for item in extracted_items:
                st.markdown(f"- {item}")

        # Store reference text
        if len(raw_text.split()) > 3500:
            st.session_state["reference_text"] = summarize_large_rfp(client, model_name=model_name, text=raw_text)
        else:
            st.session_state["reference_text"] = raw_text

    reference_text = st.session_state.get("reference_text", "")

    # ========================================
    # PREVIEW TABS (OUTSIDE INPUT CONFIGURATION)
    # ========================================
        
    if "edited_sections" in st.session_state:
        tabs = section_preview_tabs() 
        sections = st.session_state["edited_sections"]


        template_path = "Template/Integration_TemplateV3.docx"

        def load_table(key, tag):
            if key not in st.session_state:
                st.session_state[key] = extract_table_by_tag(template_path, tag)

        # Load once
        load_table("df_raci", "{{TABLE_RACI}}")
        load_table("df_activity", "{{TABLE_ACTIVITY}}")
        load_table("df_assessment", "{{TABLE_ASSESSMENT}}")
        load_table("df_staffing", "{{TABLE_STAFFING}}")
        load_table("df_resources", "{{TABLE_RESOURCES}}")
        load_table("df_commercials", "{{TABLE_COMMERCIALS}}")
        load_table("df_milestones", "{{TABLE_MILESTONES}}")
        load_table("df_interaction", "{{TABLE_INTERACTION}}")
        load_table("df_management", "{{TABLE_MANAGEMENT}}")
        load_table("df_classification", "{{TABLE_CLASSIFICATION}}")
        load_table("df_process", "{{TABLE_PROCESS}}")


            # Add tables below corresponding sections
        for i, tab in enumerate(tabs):
            with tab:
                title = sections[i]["title"]
                text = sections[i]["preview"]  # show preview version
                # st.markdown(text)

                if title == "Project Scope":
                    st.markdown("---")
                    st.subheader("📊 Section Tables")
                    with st.expander("Migration Assessment", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_assessment"],
                            key="edit_assessment"
                        )
                        if new_value is not None and not new_value.equals(st.session_state["df_assessment"]):
                            st.session_state["df_assessment"] = new_value.copy()
                
                if title == "Resource Allocation":
                    st.markdown("---")
                    st.subheader("📊 Section Tables")
                    with st.expander("Activity Table", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_activity"],
                            key="edit_activity"
                        )
                        if new_value is not None and not new_value.equals(st.session_state["df_activity"]):
                            st.session_state["df_activity"] = new_value.copy()
                    
                
                    with st.expander("Resource Allocation", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_resources"],
                            key="edit_resources"
                        )
                        if new_value is not None and not new_value.equals(st.session_state["df_resources"]):
                            st.session_state["df_resources"] = new_value.copy()


                    with st.expander("RACI Matrix", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_raci"],
                            key="edit_raci"
                        )

                        # Only update if user actually edited
                        if new_value is not None and not new_value.equals(st.session_state["df_raci"]):
                            st.session_state["df_raci"] = new_value.copy()

                if title == "Governance":
                    st.markdown("---")
                    st.subheader("📊 Section Tables")
                    with st.expander("INTERACTION Table", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_interaction"],
                            key="edit_interaction"
                        )
                        if new_value is not None and not new_value.equals(st.session_state["df_activity"]):
                            st.session_state["df_interaction"] = new_value.copy()
                    
                
                    with st.expander("Issue Management", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_management"],
                            key="edit_management"
                        )
                        if new_value is not None and not new_value.equals(st.session_state["df_resources"]):
                            st.session_state["df_management"] = new_value.copy()


                    with st.expander("Issue Classification", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_classification"],
                            key="edit_classification"
                        )

                        # Only update if user actually edited
                        if new_value is not None and not new_value.equals(st.session_state["df_raci"]):
                            st.session_state["df_classification"] = new_value.copy()

                    with st.expander("Escalation Process", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_process"],
                            key="edit_process"
                        )

                        # Only update if user actually edited
                        if new_value is not None and not new_value.equals(st.session_state["df_raci"]):
                            st.session_state["df_process"] = new_value.copy()

                
                if title == "Commercials & Payment Terms":
                    st.markdown("---")
                    st.subheader("📊 Section Tables")

                    with st.expander("Commercials", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_commercials"],
                            key="edit_com"
                        )
                        if new_value is not None and not new_value.equals(st.session_state["df_commercials"]):
                            st.session_state["df_commercials"] = new_value.copy()
                             

                    with st.expander("Payment Milestones", expanded=False):
                        new_value = st.data_editor(
                            st.session_state["df_milestones"],
                            key="edit_milestones"
                        )
                        if new_value is not None and not new_value.equals(st.session_state["df_milestones"]):
                            st.session_state["df_milestones"] = new_value.copy()
                                
    # ========================================
    # DOWNLOAD BUTTON (V2 TEMPLATE)
    # ========================================
    
    if "edited_sections" in st.session_state:
        buffer = io.BytesIO()
        
        # NOTE: You'll need to create a NEW template for V2 with updated placeholders
        template_path = "Template/Integration_TemplateV3.docx"
        
        # If V2 template doesn't exist yet, fallback to original
        if not os.path.exists(template_path):
            st.warning("⚠ V2 template not found, using original template. Please create Integration_TemplateV3.docx")
            template_path = "Template/Integration_TemplateV3.docx"
        
        final_doc = Document(template_path)

        # Basic replacements
        
        replace_submission_date(final_doc)
        doc_no = generate_document_number(client_name)
        insert_document_number(final_doc, "<DOCUMENT_NO>", doc_no)

        # V2 RESTRUCTURED placeholder map
        placeholder_map = {
            "Executive Summary": "<EXEC_SUMMARY>",
            "Why Crave": "<WHY_CRAVE>",
            "About Crave InfoTech": "<ABOUT_CRAVE>",
            "Our Understanding": "<OUR_SOL>",  # NEW placeholder
            "Project Scope": "<PROJECT_SCOPE>",
            "Solution": "<SOLUTION>",
            "Project Approach": "<DELIVERY_APPROACH>",
            "Resource Allocation": "<RESOURCE_ALLOCATION>",
            "Project Timelines": "<TIMELINES>",  
            "Commercials & Payment Terms": "<COMMERCIALS>",
            "Governance": "<GOVERNANCE>",
            "Sign Off": "<SIGN_OFF>",
            "Key Assumptions": "<KEY_ASSUMPTIONS>"
        }
        


        for sec in st.session_state["edited_sections"]:
            title = sec["title"]
            content = sec["final"]


            if title in placeholder_map:
                ph = placeholder_map[title]

                if title == "Executive Summary":
                    remove_static_before_placeholder(final_doc, ph)
                    remove_static_after_placeholder(final_doc, ph)
                    insert_formatted_text(final_doc, ph, content, replace=True)

                    # if sec.get("why_crave"):
                    #     insert_formatted_text(final_doc, "<WHY_CRAVE>", sec["why_crave"], replace=True)
                else:
                        if title == "Governance":
                            # remove_static_before_placeholder(final_doc, ph)
                            remove_static_after_placeholder(final_doc, ph)

                        insert_formatted_text(final_doc, ph, content, replace=True)

                
        # Insert PPT assets

        if "slide8_table" in st.session_state:
            insert_formatted_text(final_doc, "<ADAPTER_TABLE>", st.session_state["slide8_table"], replace=True)
        if "slide9_table" in st.session_state:
            insert_formatted_text(final_doc, "<KEY_TABLE>", st.session_state["slide9_table"], replace=True)
        if "slide7_text" in st.session_state:
            insert_formatted_text(final_doc, "<SLIDE7_TEXT>", st.session_state["slide7_text"], replace=True)
        if "slide18_resources_table" in st.session_state:
            insert_formatted_text(final_doc, "<RESOURCES_TABLE>", st.session_state["slide18_resources_table"], replace=True)
        if "total_interfaces" in st.session_state:
            replace_inline_placeholder(final_doc, "<TOTAL_INTERFACES>", st.session_state["total_interfaces"])


        # Insert assessment table at LLM placeholder
        if "df_assessment" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "[[REFER ASSESSMENT TABLE BELOW]]",
                st.session_state["df_assessment"]
            )
        remove_table_by_tag(final_doc, "{{TABLE_ASSESSMENT}}")

        # Insert activity table at LLM placeholder
        remove_table_by_tag(final_doc, "{{TABLE_ACTIVITY}}")
        if "df_activity" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "[[Refer Activity Table Below]]",
                st.session_state["df_activity"]
            )

        # Staffing & Resources: update existing template tables in-place
        if "df_staffing" in st.session_state:
            update_table_in_doc(final_doc, "{{TABLE_STAFFING}}", st.session_state["df_staffing"])

        if "df_resources" in st.session_state:
            update_table_in_doc(final_doc, "{{TABLE_RESOURCES}}", st.session_state["df_resources"])

        # ----- RACI MATRIX -----
        # Remove template RACI table and insert where LLM placed placeholder
        remove_table_by_tag(final_doc, "{{TABLE_RACI}}")
        if "df_raci" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "[[Refer RACI Table Below]]",
                st.session_state["df_raci"]
            )

        # If slide18 resources existed, keep the formatted insertion
        if "slide18_resources_table" in st.session_state:
            insert_formatted_text(final_doc, "[[Refer Resource Table Below]]", st.session_state["slide18_resources_table"], replace=True)

        # ----- GOVERNANCE TABLES -----
        remove_table_by_tag(final_doc, "{{REFER TABLE_Interaction BELOW}}")
        remove_table_by_tag(final_doc, "{{REFER TABLE_Management BELOW}}")
        remove_table_by_tag(final_doc, "{{REFER TABLE_Classification BELOW}}")
        remove_table_by_tag(final_doc, "{{REFER TABLE_Process BELOW}}")

        if "df_interaction" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "{{TABLE_INTERACTION}}",
                st.session_state["df_interaction"]
            )

        if "df_management" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "{{TABLE_MANAGEMENT}}",
                st.session_state["df_management"]
            )

        if "df_classification" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "{{TABLE_CLASSIFICATION}}",
                st.session_state["df_classification"]
            )

        if "df_process" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "{{TABLE_PROCESS}}",
                st.session_state["df_process"]
            )

                
        # # ----- COMMERCIALS TABLE -----
        remove_table_by_tag(final_doc, "{{TABLE_COMMERCIALS}}")
        if "df_commercials" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "[[Refer COMMERCIAL TABLE BELOW]]",
                st.session_state["df_commercials"]
            )

        # ----- PAYMENT TERMS TABLE -----
        remove_table_by_tag(final_doc, "{{TABLE_MILESTONES}}")
        if "df_milestones" in st.session_state:
            insert_table_at_placeholder(
                final_doc,
                "[[Refer PAYMENT TERMS TABLE BELOW]]",
                st.session_state["df_milestones"]
            )

        # 7. Insert architecture diagram from Images folder
        architecture_path = "Images/Architecture.png"

        if os.path.exists(architecture_path):
            with open(architecture_path, "rb") as img_file:
                img_bytes = img_file.read()
                insert_image_at_placeholder(final_doc, "[[ARCHITECTURE_IMG]]", img_bytes)
        else:
            st.warning("⚠ Architecture diagram not found in Images folder.")

        
        replace_client_name_in_doc(final_doc, client_name)
        
        # 8. Insert Project Approach diagram (for Section 6)
        approach_path = "Images/Approach.png"

        if os.path.exists(approach_path):
            with open(approach_path, "rb") as img_file:
                img_bytes = img_file.read()
                insert_image_at_placeholder(final_doc, "[[APPROACH_IMG]]", img_bytes)
        else:
            st.warning("⚠ Approach diagram not found in Images folder.")

        # 9. Insert Organization Structure diagram for Section 7.1
        organization_path = "Images/Organization.png"

        if os.path.exists(organization_path):
            with open(organization_path, "rb") as img_file:
                img_bytes = img_file.read()
                insert_image_at_placeholder(final_doc, "[[ORGANIZATION_IMG]]", img_bytes)
        else:
            st.warning("⚠ Organization structure diagram not found: Images/Organization.png")
        
        # 10. Insert EcoVadis rating image
        ecovadis_path = "Images/Rating.png"
        if os.path.exists(ecovadis_path):
            with open(ecovadis_path, "rb") as img_file:
                img_bytes = img_file.read()
                insert_image_at_placeholder(final_doc, "[[RATING_IMG]]", img_bytes)
        else:
            st.warning("⚠ EcoVadis rating image not found: Images/Rating.png")
                
        if "slide17_image" in st.session_state:
            insert_image_at_placeholder(final_doc, "[[PROJECT_PLAN_IMG]]", st.session_state["slide17_image"])

        # 10. Insert architecture diagram from Images folder
        lilyss_path = "Images/Lily.png"

        if os.path.exists(lilyss_path):
            with open(lilyss_path, "rb") as img_file:
                img_bytes = img_file.read()
                insert_image_at_placeholder(final_doc, "[[IMAGE:ELI_LILLY]]", img_bytes)
                insert_image_at_placeholder(final_doc, "[[IMAGE:ELILILLY]]", img_bytes)
        else:
            st.warning("⚠ Lily Success Story image not found in Images folder.")

        merckss_path = "Images/Merck.png"

        if os.path.exists(merckss_path):
            with open(merckss_path, "rb") as img_file:
                img_bytes = img_file.read()
                insert_image_at_placeholder(final_doc, "[[IMAGE:MERCK]]", img_bytes)
        else:
            st.warning("⚠ Merck Success Story image not found in Images folder.")

        final_doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Download Final SOW Document",
            data=buffer,
            file_name=f"Integration_SOW_V2_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
