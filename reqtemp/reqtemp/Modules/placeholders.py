from datetime import datetime
from docx.oxml.ns import qn

from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def replace_client_name_in_doc(doc, client_name):
    targets = [
        "(Customer Name)",
        "Customer Name",
        "(Customer name)",
        "(customer name)",
        "#CUSTNAME"
    ]

    # ----------------------------------------------------------------------
    # 1. Replace inside normal paragraphs + runs
    # ----------------------------------------------------------------------
    for p in doc.paragraphs:
        for run in p.runs:
            for t in targets:
                if t in run.text:
                    run.text = run.text.replace(t, client_name)

    # ----------------------------------------------------------------------
    # 2. Replace inside tables
    # ----------------------------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        for t in targets:
                            if t in run.text:
                                run.text = run.text.replace(t, client_name)

    # ----------------------------------------------------------------------
    # 3. Replace inside headers & footers
    # ----------------------------------------------------------------------
    for section in doc.sections:
        header = section.header
        footer = section.footer

        for p in header.paragraphs:
            for run in p.runs:
                for t in targets:
                    if t in run.text:
                        run.text = run.text.replace(t, client_name)

        for p in footer.paragraphs:
            for run in p.runs:
                for t in targets:
                    if t in run.text:
                        run.text = run.text.replace(t, client_name)

    # ----------------------------------------------------------------------
    # 4. Replace inside textboxes (cover page shapes) 🔥🔥🔥
    # ----------------------------------------------------------------------
    def replace_in_textboxes(element):
        for node in element.iter():
            if node.tag == qn("w:t"):  # Text node
                for t in targets:
                    if t in node.text:
                        node.text = node.text.replace(t, client_name)

    replace_in_textboxes(doc.element)

    return doc

from datetime import datetime

def replace_submission_date(doc):
    placeholder = "<SUBMISSION_DATE>"
    today = datetime.now().strftime("%d-%m-%y")   # format: 18-11-25

    # paragraphs
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, today)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, today)

    # headers & footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, today)
        for p in section.footer.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, today)


def replace_inline_placeholder(doc, placeholder, value):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, value)



def generate_document_number(client_name, version=1):
    """
    Generates document number like: POM/25001CPA
    No region code included.
    """
    import datetime

    prefix = "POM"   # PIPO Migration

    # Year in 2 digits
    year = str(datetime.datetime.now().year)[-2:]

    # Version padded to 3 digits
    version_str = str(version).zfill(3)

    # Extract client initials → first 3 alphabetic characters of each word
    client_initials = ""
    for word in client_name.split():
        for ch in word:
            if ch.isalpha():
                client_initials += ch.upper()
                break

    client_initials = client_initials[:3] if client_initials else "CLT"

    # Combine
    return f"{prefix}/{year}{version_str}{client_initials}"

def insert_document_number(doc, placeholder, value):
    # paragraphs
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, value)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, value)

    # headers & footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, value)
        for p in section.footer.paragraphs:
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, value)
