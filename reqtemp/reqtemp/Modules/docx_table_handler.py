import pandas as pd
from docx import Document

def extract_table_by_tag(doc_path, tag):
    """
    Finds a table containing a specific tag in the first cell (0,0),
    returns data as a Pandas DataFrame.
    """
    try:
        doc = Document(doc_path)
    except Exception as e:
        return pd.DataFrame()

    for table in doc.tables:
        try:
            # Check first cell for the unique tag
            # We strip whitespace to match cleanly
            first_cell_text = table.cell(0, 0).text.strip()
            
            if tag in first_cell_text:
                data = []
                keys = None
                
                # Assume Row 0 is the Header.
                # We remove the tag from the header name for the DataFrame column
                keys = [cell.text.replace(tag, "").strip() for cell in table.rows[0].cells]
                
                # Iterate from Row 1 onwards for data
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)
                
                # Create DataFrame
                if keys and len(keys) == len(data[0] if data else []):
                     df = pd.DataFrame(data, columns=keys)
                elif data:
                     # Fallback if column count mismatch
                     df = pd.DataFrame(data)
                else:
                     df = pd.DataFrame(columns=keys)
                     
                return df
                
        except (IndexError, AttributeError):
            continue
            
    return pd.DataFrame() # Return empty if not found


def update_table_in_doc(doc, tag, modified_df):
    """
    Finds the table by tag in the doc object and updates its cells 
    with values from the modified DataFrame.
    """
    for table in doc.tables:
        try:
            first_cell_text = table.cell(0, 0).text
            if tag in first_cell_text:
                
                # 1. Update existing rows
                # We skip row 0 because it's the header in the Word Doc
                for i, row in enumerate(modified_df.itertuples(index=False)):
                    
                    # If DataFrame has more rows than the Word table (user added rows), add rows to Word table
                    # i + 1 because we are skipping the header row (0)
                    if i + 1 >= len(table.rows):
                        table.add_row()
                    
                    docx_row = table.rows[i + 1] 
                    
                    for j, value in enumerate(row):
                        # Safety check: ensure we don't try to write to a column that doesn't exist in Word
                        if j < len(docx_row.cells):
                            # Set text. Note: This replaces cell content but keeps cell formatting (borders/bg)
                            # to some extent, though run-level styling (bold inside cell) might be reset to default.
                            docx_row.cells[j].text = str(value)
                            
                # 2. Cleanup: Remove the Tag from the header cell for the final print
                header_cell = table.cell(0,0)
                if tag in header_cell.text:
                    header_cell.text = header_cell.text.replace(tag, "").strip()
                
                return doc
        except IndexError:
            continue
            
    return doc