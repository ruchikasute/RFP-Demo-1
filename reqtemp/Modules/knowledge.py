# import os
# from docx import Document
# from PyPDF2 import PdfReader

# def load_knowledge_text():
#     """
#     Loads all reference SOWs from /knowledge folder,
#     extracts only clean text (no tables, no weird formatting),
#     and merges them into a single style guide.
#     """

#     folder = "Knowledge_Repo/Integration"
#     if not os.path.exists(folder):
#         return ""

#     knowledge_texts = []

#     for file in os.listdir(folder):
#         path = os.path.join(folder, file)

#         try:
#             if file.lower().endswith(".docx"):
#                 doc = Document(path)
#                 # Only paragraphs with meaningful text
#                 txt = "\n".join(
#                     p.text.strip()
#                     for p in doc.paragraphs
#                     if len(p.text.strip()) > 20  # avoid junk
#                 )
#                 knowledge_texts.append(txt)

#             elif file.lower().endswith(".pdf"):
#                 pdf = PdfReader(path)
#                 pages = []
#                 for p in pdf.pages:
#                     try:
#                         t = p.extract_text()
#                         if t and len(t.strip()) > 20:
#                             pages.append(t)
#                     except:
#                         continue
#                 knowledge_texts.append("\n".join(pages))

#         except Exception as e:
#             print("Failed to read knowledge doc:", file, e)

#     # Merge & clean
#     merged = "\n\n".join(knowledge_texts)
#     return merged.strip()


import os
import chromadb
from sentence_transformers import SentenceTransformer
from docx import Document
from PyPDF2 import PdfReader

# -----------------------------
# MODEL + VECTOR DB SETTINGS
# -----------------------------
EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
DB_DIR = "vector_db"


def get_vector_client():
    return chromadb.PersistentClient(path="vector_db")


# -----------------------------
# 1. LOAD TEXT FROM FILES
# -----------------------------
def load_text_from_file(path):
    text = ""

    if path.lower().endswith(".docx"):
        doc = Document(path)
        text = "\n".join(
            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()
        )

    elif path.lower().endswith(".pdf"):
        pdf = PdfReader(path)
        pages = []
        for p in pdf.pages:
            try:
                t = p.extract_text()
                if t and len(t.strip()) > 20:
                    pages.append(t)
            except:
                continue
        text = "\n".join(pages)

    return text.strip()


# -----------------------------
# 2. SPLIT TEXT INTO CHUNKS
# -----------------------------
def chunk_text(text, chunk_size=600):
    words = text.split()
    chunks = []
    current = []

    for w in words:
        current.append(w)
        if len(current) >= chunk_size:
            chunks.append(" ".join(current))
            current = []

    if current:
        chunks.append(" ".join(current))

    return chunks


# -----------------------------
# 3. BUILD / UPDATE VECTOR DB
# -----------------------------
def build_vector_store(folder_path , category_name):
    client = get_vector_client()

    # Create or load collection
    collection = client.get_or_create_collection(
        name="knowledge_repo",
        metadata={"hnsw:space": "cosine"},
    )

    for file in os.listdir(folder_path):
        path = os.path.join(folder_path, file)
        if not os.path.isfile(path):
            continue

        try:
            text = load_text_from_file(path)
            if not text or len(text) < 200:
                continue

            chunks = chunk_text(text)
            embeddings = EMBED_MODEL.encode(chunks).tolist()
            ids = [f"{file}_{i}" for i in range(len(chunks))]

            collection.add(
                documents=chunks,
                embeddings=embeddings,
                ids=ids,
                metadatas=[{"category": category_name}] * len(chunks)
)

        except Exception as e:
            print("Failed loading:", file, e)

    print("🔥 Vector Database Updated")
    return True


# -----------------------------
# 4. SEARCH MOST RELEVANT CHUNKS
# -----------------------------
def similarity_search(query, category=None,top_k=3):
    client = get_vector_client()

    try:
        collection = client.get_collection("knowledge_repo")
    except:
        return []

    q_emb = EMBED_MODEL.encode([query]).tolist()

    results = collection.query(
        query_embeddings=q_emb,
        n_results=top_k,
        where={"category": category} if category else {}
    )

    return results["documents"][0]


# -----------------------------
# 5. MAIN FUNCTION USED BY APP
# -----------------------------
def load_knowledge_text(query="SAP migration SOW style", category=None):
    """
    Returns only the most relevant 5–7 repo chunks,
    NOT the full repository → keeps LLM fast & focused.
    """
    try:
        chunks = similarity_search(query,category=category, top_k=7)
        if chunks:
            return "\n\n".join(chunks)
    except:
        pass

    return ""


# -----------------------------
# 6. ONE-TIME INIT FUNCTION
# -----------------------------
def initialize_vector_db():
    base_folder = "Knowledge_Repo"
    client = get_vector_client()

    for folder in os.listdir(base_folder):
        folder_path = os.path.join(base_folder, folder)
        if not os.path.isdir(folder_path):
            continue

        print(f"🔍 Indexing folder: {folder}")

        # Build store for a specific folder
        build_vector_store(folder_path, folder)
